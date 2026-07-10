"""
EBE-FRM-006 — Pauta de Avaliação dos Alunos
Modelo oficial de lançamento de notas por turma/curso/módulo.
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line


def paisagem(doc):
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)


def cell_text(cell, text, bold=False, color=None, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FONTE_CORPO
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def header_row(row, labels, size=8):
    for cell, label in zip(row.cells, labels):
        _shade_cell(cell, HEX_PRIMARIA)
        cell_text(cell, label, bold=True, color=RGBColor(255,255,255), size=size)


def gerar():
    doc = novo_documento("Pauta de Avaliação", "EBE-FRM-006")
    paisagem(doc)

    add_capa(doc,
        supratitulo="Formulário Académico N.º 6",
        titulo="Pauta de Avaliação dos Alunos",
        subtitulo="Registo oficial de notas, frequência, situação académica e validação final",
        codigo="EBE-FRM-006", ano="2026")
    add_marco_filosofico(doc)

    h1(doc, "Identificação da Turma / Unidade Avaliada", numero=1)
    tbl = doc.add_table(rows=6, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    campos = [
        ("Instituto", "«[Nome do Instituto]»", "Escola", "«[Nome da Escola]»"),
        ("Curso", "«[Nome do Curso]»", "Módulo", "«[Nome/N.º do Módulo]»"),
        ("Apostila / Unidade", "«[Tema ou unidade avaliada]»", "Nível", "«[Discípulo/Crescimento/Servir/Multiplicação]»"),
        ("Docente", "«[Nome do(a) docente]»", "Turma", "«[Código da turma]»"),
        ("Período", "«[mês/ano]»", "Carga horária", "«[___ h]»"),
        ("Data de emissão", "«[__/__/____]»", "Média mínima", "6,0"),
    ]
    for r, linha in enumerate(campos):
        for c, txt in enumerate(linha):
            cell = tbl.rows[r].cells[c]
            if c in (0,2):
                _shade_cell(cell, "E8F1EC")
                cell_text(cell, txt, bold=True, color=COR_PRIMARIA, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                cell_text(cell, txt, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    h1(doc, "Pauta Geral — Avaliação Integral", numero=2)
    paragrafo(doc, "A avaliação institucional da EBE considera as quatro dimensões formativas: Conhecer (40 %), Crer (20 %), Viver (20 %) e Servir (20 %). A média final é ponderada e a aprovação exige média mínima 6,0 e frequência mínima institucional de 75 %.")
    citacao(doc, "Tem cuidado de ti mesmo e da doutrina. Persevera nestes deveres; porque, fazendo isto, te salvarás, tanto a ti mesmo como aos que te ouvem.", "1 Timóteo 4.16")

    labels = [
        "N.º", "Matrícula", "Nome completo do(a) aluno(a)", "Freq. %",
        "Conhecer\n40 %", "Crer\n20 %", "Viver\n20 %", "Servir\n20 %",
        "Média\nfinal", "Sit.", "Obs. / Recuperação"
    ]
    tbl = doc.add_table(rows=1, cols=len(labels))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], labels, size=7.5)
    for i in range(1, 31):
        row = tbl.add_row().cells
        vals = [str(i), "", "«[Nome do(a) aluno(a)]»" if i == 1 else "", "", "", "", "", "", "", "", ""]
        for idx, val in enumerate(vals):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx in (2,10) else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(row[idx], val, size=7.5, align=align)

    h1(doc, "Instrumentos de Avaliação Utilizados", numero=3)
    paragrafo(doc, "Registe aqui os instrumentos concretos aplicados no módulo/curso e o respectivo peso interno dentro de cada dimensão.")
    labels = ["N.º", "Instrumento", "Dimensão", "Peso interno", "Data", "Descrição / Observações"]
    tbl = doc.add_table(rows=1, cols=len(labels))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], labels)
    exemplos = [
        ("1", "Prova escrita", "Conhecer", "«[__ %]»", "«[__/__/____]»", "Avaliação de conceitos, textos bíblicos e doutrina."),
        ("2", "Trabalho escrito", "Conhecer/Crer", "«[__ %]»", "«[__/__/____]»", "Pesquisa, argumentação e uso bíblico-teológico."),
        ("3", "Participação e conduta", "Viver", "«[__ %]»", "«[__/__/____]»", "Pontualidade, responsabilidade, vida devocional e respeito."),
        ("4", "Prática ministerial", "Servir", "«[__ %]»", "«[__/__/____]»", "Serviço, discipulado, evangelização ou projecto prático."),
        ("5", "Outro", "«[Dimensão]»", "«[__ %]»", "«[__/__/____]»", "«[Descrição]»"),
    ]
    for data in exemplos:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=8, align=WD_ALIGN_PARAGRAPH.LEFT if i in (1,5) else WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)

    h1(doc, "Pauta Detalhada — Lançamento por Instrumento", numero=4)
    paragrafo(doc, "Use esta página quando desejar lançar mais de uma nota por dimensão antes de consolidar a média final na pauta geral.")
    labels = [
        "N.º", "Nome", "C1", "C2", "C3", "Média C", "CR1", "CR2", "Média CR", "V1", "V2", "Média V", "S1", "S2", "Média S", "Média Final"
    ]
    tbl = doc.add_table(rows=1, cols=len(labels))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], labels, size=7)
    for i in range(1, 26):
        row = tbl.add_row().cells
        vals = [str(i), ""] + [""]*(len(labels)-2)
        for idx, val in enumerate(vals):
            cell_text(row[idx], val, size=7, align=WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER)

    h1(doc, "Legenda Académica", numero=5)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Código", "Situação", "Critério", "Observação"])
    legenda = [
        ("A", "Aprovado", "Média final ≥ 6,0 e frequência ≥ 75 %", "Certificável."),
        ("AC", "Aprovado por compensação", "Critério excepcional da Coordenação", "Deve haver registo formal."),
        ("R", "Reprovado", "Média < 6,0 e/ou frequência insuficiente", "Pode requerer recuperação, se aplicável."),
        ("EC", "Em curso", "Avaliação ainda não encerrada", "Usar antes da pauta final."),
        ("DI", "Dispensado por equivalência", "Equivalência aprovada", "Anexar comprovativo."),
        ("TR", "Trancado", "Matrícula suspensa conforme pedido/decisão", "Sem nota final."),
    ]
    for data in legenda:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)

    h1(doc, "Fórmulas Institucionais", numero=6)
    paragrafo(doc, "Média Final = (Conhecer × 0,40) + (Crer × 0,20) + (Viver × 0,20) + (Servir × 0,20).")
    paragrafo(doc, "Frequência (%) = (presenças válidas ÷ total de aulas/actividades previstas) × 100.")
    paragrafo(doc, "A aprovação ordinária exige simultaneamente: média final mínima de 6,0 e frequência mínima de 75 %.")

    h1(doc, "Escala de Notas", numero=7)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Intervalo", "Menção", "Descrição"])
    escala = [
        ("9,0 – 10,0", "Excelente", "Domínio elevado, maturidade e aplicação consistente."),
        ("7,5 – 8,9", "Muito bom", "Bom domínio, boa aplicação e participação regular."),
        ("6,0 – 7,4", "Satisfatório", "Cumpre os requisitos mínimos com aproveitamento."),
        ("5,0 – 5,9", "Suficiente", "Aproveitamento parcial, abaixo da média institucional."),
        ("0,0 – 4,9", "Insuficiente", "Não cumpre os objectivos mínimos da unidade."),
    ]
    for data in escala:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT)

    h1(doc, "Validação da Pauta", numero=8)
    paragrafo(doc, "A pauta somente se torna oficial após conferência da Secretaria Acadêmica e validação da Coordenação Acadêmica. Rasuras devem ser evitadas; eventuais correcções devem ser rubricadas.")
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, label in zip(tbl.rows[0].cells, ["Docente", "Secretaria Acadêmica", "Coordenação Acadêmica", "Data de homologação"]):
        _shade_cell(cell, "E8F1EC")
        cell_text(cell, label, bold=True, color=COR_PRIMARIA, size=8.5)
    vals = ["\n\n______________________________\nAssinatura", "\n\n______________________________\nAssinatura", "\n\n______________________________\nAssinatura", "\n\n«[__/__/____]»"]
    for cell, val in zip(tbl.rows[1].cells, vals):
        cell_text(cell, val, size=8.5)

    selo_final(doc)
    out = os.path.join(os.path.dirname(__file__), "EBE-FRM-006_Pauta_de_Avaliacao_dos_Alunos.docx")
    doc.save(out)
    print("OK:", out)

if __name__ == "__main__":
    gerar()
