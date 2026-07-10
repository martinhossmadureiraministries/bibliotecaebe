"""
EBE-MOD-TCC-001 — Modelo de TCC / Monografia Final
Modelo universal para todos os institutos, escolas e cursos da EBE.
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line


def cell_text(cell, text, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FONTE_CORPO
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def header_row(row, labels):
    for cell, label in zip(row.cells, labels):
        _shade_cell(cell, HEX_PRIMARIA)
        cell_text(cell, label, bold=True, color=RGBColor(255,255,255), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def quadro(doc, titulo, texto):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "E8F1EC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(titulo + " — ")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA
    r2 = p.add_run(texto)
    r2.font.name = FONTE_CORPO
    r2.font.size = Pt(11)


def gerar():
    doc = novo_documento("Modelo de TCC / Monografia Final", "EBE-MOD-TCC-001")

    add_capa(doc,
        supratitulo="Modelo Pedagógico Oficial",
        titulo="Modelo de TCC / Monografia Final",
        subtitulo="Trabalho de Conclusão aplicável a todos os Institutos, Escolas e Cursos da Escola Bíblica Epignósis",
        codigo="EBE-MOD-TCC-001", ano="2026")
    add_marco_filosofico(doc)

    h1(doc, "Natureza e Finalidade do Trabalho Final", numero=1)
    paragrafo(doc, "O Trabalho de Conclusão de Curso (TCC) ou Monografia Final é o instrumento académico por meio do qual o(a) aluno(a) demonstra maturidade bíblica, teológica, espiritual, ministerial e metodológica ao final de um curso, escola, instituto, programa ou da Formação Completa Epignósis.")
    paragrafo(doc, "Este modelo é universal e pode ser utilizado por todos os Institutos e Cursos da Escola Bíblica Epignósis, com as devidas adaptações ao nível formativo, à carga horária e à natureza do curso.")
    citacao(doc, "Procura apresentar-te a Deus aprovado, como obreiro que não tem de que se envergonhar, que maneja bem a palavra da verdade.", "2 Timóteo 2.15")
    quadro(doc, "Princípio Epignósis", "O TCC não é apenas uma exigência académica. É um acto de serviço ao Reino, uma oportunidade de transformar estudo em edificação da Igreja e da sociedade.")

    h1(doc, "Modalidades Aceites", numero=2)
    paragrafo(doc, "Conforme a natureza do curso e a orientação da Coordenação, o trabalho final poderá assumir uma das seguintes modalidades:")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Modalidade", "Descrição", "Mais indicada para"])
    for data in [
        ("Monografia bíblico-teológica", "Pesquisa académica escrita, com problema, objectivos, fundamentação bíblica, revisão bibliográfica, análise e conclusão.", "Ciências Bíblicas, Ciências Teológicas, Pesquisa Bíblica."),
        ("Projecto de intervenção ministerial", "Plano aplicado a uma igreja, célula, ministério, missão, escola bíblica, grupo de discipulado ou realidade comunitária.", "Ministerial, Cinco Ministérios, Liderança, Missões."),
        ("Estudo bíblico-exegético", "Análise aprofundada de um texto, tema ou livro bíblico, respeitando contexto, língua, género literário e aplicação.", "Hermenêutica, Exegese, Pregação, Ciências Bíblicas."),
        ("Relatório reflexivo de prática", "Descrição, análise e avaliação espiritual/pedagógica de uma experiência ministerial supervisionada.", "Formação Espiritual, Serviço, Estágio, Missões."),
        ("Produto pedagógico", "Criação de apostila, plano de curso, série de estudos, manual, devocional ou recurso formativo, acompanhado de fundamentação teórica.", "Docência, Educação Cristã, Discipulado, Liderança."),
    ]:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=9)

    h1(doc, "Identificação Institucional do Trabalho", numero=3)
    tbl = doc.add_table(rows=8, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fields = [
        ("Instituto", "«[Nome do Instituto]»"),
        ("Escola", "«[Nome da Escola]»"),
        ("Curso", "«[Nome do Curso]»"),
        ("Nível formativo", "«[Discípulo / Crescimento / Servir / Multiplicação]»"),
        ("Modalidade do trabalho", "«[Monografia / Projecto / Estudo Exegético / Relatório / Produto Pedagógico]»"),
        ("Aluno(a)", "«[Nome completo do(a) aluno(a)]»"),
        ("Orientador(a)", "«[Nome do(a) orientador(a)]»"),
        ("Local e ano", "«[Luanda]», «[2026]»"),
    ]
    for i, (a,b) in enumerate(fields):
        _shade_cell(tbl.rows[i].cells[0], "E8F1EC")
        cell_text(tbl.rows[i].cells[0], a, bold=True, color=COR_PRIMARIA, size=10)
        cell_text(tbl.rows[i].cells[1], b, size=10)

    h1(doc, "Institutos e Possíveis Linhas de Pesquisa", numero=4)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Instituto", "Exemplos de linhas de pesquisa / aplicação"])
    linhas = [
        ("Formação Cristã", "Discipulado, novo nascimento, identidade em Cristo, santidade, oração, fundamentos da fé."),
        ("Ciências Bíblicas", "Introdução bíblica, panorama bíblico, exegese, hermenêutica, livros bíblicos, géneros literários."),
        ("Ciências Teológicas", "Doutrina de Deus, cristologia, pneumatologia, soteriologia, eclesiologia, escatologia."),
        ("Formação Espiritual", "Vida devocional, disciplinas espirituais, carácter cristão, maturidade, cura interior bíblica."),
        ("Ministerial", "Pregação, ensino, aconselhamento bíblico, capelania, administração ministerial."),
        ("Reino e Poder", "Reino de Deus, autoridade espiritual, dons espirituais, oração por enfermos, libertação bíblica."),
        ("Cinco Ministérios", "Apóstolos, profetas, evangelistas, pastores e mestres na edificação da Igreja."),
        ("Missões", "Evangelização, plantação de igrejas, missões urbanas, transculturais, mobilização missionária."),
        ("Liderança e Multiplicação", "Liderança servidora, mentoria, células, formação de líderes, gestão e visão ministerial."),
        ("Pesquisa Bíblica e Excelência", "Metodologia, escrita académica, produção de materiais, investigação teológica aplicada."),
    ]
    for inst, linha in linhas:
        row = tbl.add_row().cells
        cell_text(row[0], inst, bold=True, color=COR_PRIMARIA, size=9)
        cell_text(row[1], linha, size=9)

    page_break(doc)

    h1(doc, "Estrutura Obrigatória do Trabalho", numero=5)
    paragrafo(doc, "Salvo orientação específica da Coordenação, o trabalho final deve conter os seguintes elementos:")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Parte", "Elemento", "Obrigatoriedade"])
    estrutura = [
        ("Pré-textual", "Capa", "Obrigatório"),
        ("Pré-textual", "Folha de rosto", "Obrigatório"),
        ("Pré-textual", "Declaração de originalidade", "Obrigatório"),
        ("Pré-textual", "Dedicatória", "Opcional"),
        ("Pré-textual", "Agradecimentos", "Opcional"),
        ("Pré-textual", "Resumo", "Obrigatório"),
        ("Pré-textual", "Palavras-chave", "Obrigatório"),
        ("Pré-textual", "Sumário", "Obrigatório"),
        ("Textual", "Introdução", "Obrigatório"),
        ("Textual", "Capítulo 1 — Fundamentação bíblica e teológica", "Obrigatório"),
        ("Textual", "Capítulo 2 — Revisão bibliográfica / enquadramento teórico", "Obrigatório"),
        ("Textual", "Capítulo 3 — Análise, proposta, intervenção ou estudo de caso", "Obrigatório"),
        ("Textual", "Conclusão", "Obrigatório"),
        ("Pós-textual", "Referências bibliográficas", "Obrigatório"),
        ("Pós-textual", "Anexos ou apêndices", "Opcional"),
    ]
    for data in estrutura:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=9)

    h1(doc, "Normas de Formatação", numero=6)
    lista(doc, [
        "Formato: A4, orientação vertical.",
        "Fonte: Garamond, tamanho 12 no corpo do texto.",
        "Espaçamento: 1,5 entre linhas; 6 pt depois de cada parágrafo.",
        "Margens académicas: esquerda 3 cm; direita 2,5 cm; superior 2,5 cm; inferior 2,5 cm.",
        "Texto justificado, com linguagem formal, clara e pastoralmente respeitosa.",
        "Citações bíblicas: Almeida Revista e Corrigida (ARC), entre aspas, com referência no formato: (Livro capítulo.versículo, ARC).",
        "Citações longas: recuadas, tamanho 11, itálico ou corpo diferenciado, com referência completa.",
        "Notas de rodapé: apenas quando necessárias para esclarecimento técnico.",
        "Extensão recomendada: 15–25 páginas para curso; 25–40 páginas para escola/instituto; 40–60 páginas para formação completa, salvo orientação específica.",
    ])

    h1(doc, "Capa — Modelo Editável", numero=7)
    paragrafo(doc, "Substitua os campos entre guillemets pelos dados reais do trabalho.")
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "FFFFFF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold, color in [
        ("ESCOLA BÍBLICA EPIGNÓSIS", 14, True, COR_PRIMARIA),
        ("«[Instituto]»", 12, True, COR_SECUNDARIA),
        ("«[Escola]» · «[Curso]»", 12, False, COR_TEXTO),
        ("\n\n«[TÍTULO DO TRABALHO]»", 18, True, COR_PRIMARIA),
        ("«[Subtítulo, se houver]»", 12, False, COR_TEXTO),
        ("\n\n«[Nome completo do(a) aluno(a)]»", 12, True, COR_TEXTO),
        ("\n\nTrabalho de Conclusão apresentado à Escola Bíblica Epignósis como requisito parcial para conclusão de «[Curso/Escola/Instituto/Programa]».", 11, False, COR_TEXTO),
        ("\nOrientador(a): «[Nome do(a) orientador(a)]»", 11, False, COR_TEXTO),
        ("\n\n«[Local]»\n«[Ano]»", 12, True, COR_PRIMARIA),
    ]:
        run = p.add_run(text + "\n")
        run.font.name = FONTE_TITULO
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color

    page_break(doc)

    h1(doc, "Elementos Pré-Textuais — Modelos", numero=8)
    h2(doc, "Declaração de Originalidade")
    paragrafo(doc, "Declaro, para os devidos efeitos, que o presente trabalho é de minha autoria, que todas as fontes consultadas foram devidamente identificadas e que não contém plágio, fraude académica ou apropriação indevida de ideias, textos ou produções de terceiros.")
    paragrafo(doc, "«[Local]», «[data]»")
    paragrafo(doc, "__________________________________\n«[Nome completo do(a) aluno(a)]»")

    h2(doc, "Resumo")
    paragrafo(doc, "«[Escreva aqui um resumo de 150 a 250 palavras, apresentando tema, problema, objectivo, metodologia, principais resultados e contribuição para a Igreja/Reino de Deus.]»")
    paragrafo(doc, "Palavras-chave: «[palavra 1]»; «[palavra 2]»; «[palavra 3]»; «[palavra 4]». ", bold=True)

    h1(doc, "Texto Principal — Modelo de Capítulos", numero=9)
    h2(doc, "Introdução")
    lista(doc, [
        "Apresentação do tema.",
        "Delimitação do assunto.",
        "Problema de pesquisa ou necessidade ministerial.",
        "Justificativa bíblica, teológica e prática.",
        "Objectivo geral e objectivos específicos.",
        "Metodologia utilizada.",
        "Estrutura dos capítulos.",
    ])
    paragrafo(doc, "«[Escreva aqui a introdução do trabalho.]»")

    h2(doc, "Capítulo 1 — Fundamentação Bíblica e Teológica")
    paragrafo(doc, "«[Exponha os principais textos bíblicos, doutrinas e princípios teológicos que sustentam o tema. Use a ARC como versão de referência e respeite o contexto imediato, histórico, cultural, gramatical e literário.]»")
    citacao(doc, "Toda Escritura divinamente inspirada é proveitosa para ensinar, para redarguir, para corrigir, para instruir em justiça.", "2 Timóteo 3.16")

    h2(doc, "Capítulo 2 — Revisão Bibliográfica / Enquadramento Teórico")
    paragrafo(doc, "«[Apresente autores, livros, artigos, apostilas e documentos institucionais que dialogam com o tema. Compare posições com respeito, clareza e fidelidade bíblica.]»")

    h2(doc, "Capítulo 3 — Análise, Proposta, Intervenção ou Estudo de Caso")
    paragrafo(doc, "«[Desenvolva a análise principal. Se for projecto ministerial, descreva diagnóstico, público-alvo, objectivos, plano de acção, recursos, cronograma e indicadores de fruto. Se for monografia, apresente a argumentação principal e os resultados da pesquisa.]»")

    h2(doc, "Conclusão")
    lista(doc, [
        "Retome o problema ou objectivo inicial.",
        "Resuma os principais achados.",
        "Indique implicações para a vida cristã, a Igreja e o Reino.",
        "Apresente limitações do trabalho.",
        "Sugira novas pesquisas ou acções ministeriais futuras.",
    ])
    paragrafo(doc, "«[Escreva aqui a conclusão.]»")

    h1(doc, "Referências Bibliográficas — Modelo", numero=10)
    paragrafo(doc, "As referências devem incluir todas as obras citadas no texto. Use um padrão único em todo o trabalho.")
    lista(doc, [
        "BÍBLIA SAGRADA. Almeida Revista e Corrigida. «[Editora]», «[Ano]».",
        "SOBRENOME, Nome. Título da obra em itálico. Cidade: Editora, Ano.",
        "SOBRENOME, Nome. “Título do capítulo”. In: SOBRENOME, Nome (org.). Título da obra. Cidade: Editora, Ano, p. xx–yy.",
        "AUTOR/INSTITUIÇÃO. Título do documento. Disponível em: «[URL]». Acesso em: «[data]».",
        "ESCOLA BÍBLICA EPIGNÓSIS. «[Título da apostila/documento]». Luanda: EBE, 2026.",
    ])

    h1(doc, "Cronograma Sugerido de Produção", numero=11)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Etapa", "Actividade", "Prazo sugerido", "Responsável"])
    for data in [
        ("1", "Escolha do tema e linha de pesquisa", "Semana 1", "Aluno(a) + Orientador(a)"),
        ("2", "Entrega do pré-projecto", "Semana 2", "Aluno(a)"),
        ("3", "Aprovação do tema", "Semana 3", "Coordenação/Orientador(a)"),
        ("4", "Levantamento bibliográfico", "Semanas 3–4", "Aluno(a)"),
        ("5", "Redacção da introdução e capítulo 1", "Semanas 5–6", "Aluno(a)"),
        ("6", "Redacção dos capítulos 2 e 3", "Semanas 7–9", "Aluno(a)"),
        ("7", "Conclusão e referências", "Semana 10", "Aluno(a)"),
        ("8", "Revisão com orientador(a)", "Semana 11", "Orientador(a)"),
        ("9", "Entrega final", "Semana 12", "Aluno(a)"),
        ("10", "Apresentação/defesa, se aplicável", "Semana 13", "Banca/Coordenação"),
    ]:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0,2) else WD_ALIGN_PARAGRAPH.LEFT)

    h1(doc, "Critérios de Avaliação do TCC / Monografia", numero=12)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Dimensão", "Peso", "Critérios", "Nota"])
    avaliacao = [
        ("Conhecer", "40 %", "Domínio bíblico-teológico; clareza conceitual; uso correcto das Escrituras; qualidade da pesquisa.", "«[0–10]»"),
        ("Crer", "20 %", "Coerência com a Declaração de Fé; convicções bíblicas; defesa da fé com mansidão.", "«[0–10]»"),
        ("Viver", "20 %", "Aplicação à vida; maturidade espiritual; integridade; postura académica e cristã.", "«[0–10]»"),
        ("Servir", "20 %", "Contribuição prática para a Igreja/Reino; relevância ministerial; possibilidade de multiplicação.", "«[0–10]»"),
    ]
    for data in avaliacao:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (1,3) else WD_ALIGN_PARAGRAPH.LEFT)
    paragrafo(doc, "Média final ponderada = (Conhecer × 0,40) + (Crer × 0,20) + (Viver × 0,20) + (Servir × 0,20). Média mínima institucional: 6,0.", italic=True)

    h1(doc, "Ficha de Aprovação", numero=13)
    paragrafo(doc, "O presente trabalho, intitulado «[Título]», elaborado por «[Nome do(a) aluno(a)]», foi avaliado e considerado:")
    paragrafo(doc, "Resultado:  □ Aprovado(a)   □ Aprovado(a) com correcções   □ Reprovado(a)")
    paragrafo(doc, "Nota final: «[0–10]»    Data: «[__/__/____]»")
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, label in zip(tbl.rows[0].cells, ["Orientador(a)", "Avaliador(a) 1", "Avaliador(a) 2 / Coordenação"]):
        _shade_cell(cell, "E8F1EC")
        cell_text(cell, label, bold=True, color=COR_PRIMARIA, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for cell in tbl.rows[1].cells:
        cell_text(cell, "\n\n__________________________________\nNome, assinatura e data", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    selo_final(doc)
    out = os.path.join(os.path.dirname(__file__), "EBE-MOD-TCC-001_Modelo_de_TCC_Monografia_Final.docx")
    doc.save(out)
    print("OK:", out)

if __name__ == "__main__":
    gerar()
