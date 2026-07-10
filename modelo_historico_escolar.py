"""
MODELO OFICIAL DE HISTÓRICO ESCOLAR — Escola Bíblica Epignósis

Documento académico oficial com:
  - Identificação do aluno
  - Resumo formativo por nível
  - Registo detalhado de cursos cumpridos (por instituto/escola/curso)
  - Médias, frequência, situação académica
  - Quadro de equivalências
  - Atestação institucional
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line

from docx.enum.section import WD_ORIENT


def _campo_inline(tbl, rotulo, valor=""):
    row = tbl.add_row().cells
    row[0].text = rotulo; row[1].text = valor
    _shade_cell(row[0], "E8F1EC")
    for p in row[0].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
            r.font.color.rgb = COR_PRIMARIA
    for p in row[1].paragraphs:
        for r in p.runs:
            r.font.name = FONTE_CORPO; r.font.size = Pt(10)


def _tabela_registo(doc, headers, n_linhas_vazias=8, widths=None):
    """Cria tabela de registo acadêmico com cabeçalho azul + linhas em branco."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].text = ""
        if widths:
            hdr[i].width = Cm(widths[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for _ in range(n_linhas_vazias):
        row = tbl.add_row().cells
        for i, c in enumerate(row):
            if widths:
                c.width = Cm(widths[i])
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return tbl


def _secao_form(doc, titulo_secao):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"  {titulo_secao.upper()}")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEX_PRIMARIA)
    pPr.append(shd)


def gerar():
    doc = novo_documento("Histórico Escolar", "EBE-HIS-001")

    # Mudar a 1ª secção para paisagem para acomodar tabelas largas
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ====== CABEÇALHO ======
    # Linha com logo à esquerda + título à direita
    tbl_cab = doc.add_table(rows=1, cols=2)
    tbl_cab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cab.rows[0].cells[0].width = Cm(5)
    tbl_cab.rows[0].cells[1].width = Cm(21)

    c_logo = tbl_cab.rows[0].cells[0]
    p = c_logo.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(LOGO_PATH, width=Cm(3.5))

    c_tit = tbl_cab.rows[0].cells[1]
    p = c_tit.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = c_tit.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ἐπίγνωσις  ·  Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(10)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = c_tit.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("HISTÓRICO ESCOLAR")
    r.font.name = FONTE_TITULO; r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = c_tit.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Documento académico oficial · Código EBE-HIS-«[ano]»-«[n.º]»")
    r.font.name = FONTE_CORPO; r.font.size = Pt(9)
    r.font.italic = True; r.font.color.rgb = COR_CITACAO

    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    # ====== 1. IDENTIFICAÇÃO DO ALUNO ======
    _secao_form(doc, "1.  Identificação do(a) aluno(a)")

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 4 colunas: rótulo | valor | rótulo | valor
    def _campo4(rot1, val1, rot2, val2):
        row = tbl.add_row().cells
        row[0].text = rot1; row[1].text = val1
        row[2].text = rot2; row[3].text = val2
        _shade_cell(row[0], "E8F1EC")
        _shade_cell(row[2], "E8F1EC")
        for idx in (0, 2):
            for p in row[idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(9)
                    r.font.color.rgb = COR_PRIMARIA
        for idx in (1, 3):
            for p in row[idx].paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(9)

    _campo4("Nome completo", "«[…]»", "N.º de Matrícula", "«[…]»")
    _campo4("Data de nascimento", "«[__/__/____]»", "Documento de identificação", "«[…]»")
    _campo4("Nacionalidade", "«[…]»", "Naturalidade", "«[…]»")
    _campo4("Igreja local", "«[…]»", "Denominação", "«[…]»")
    _campo4("Pastor responsável", "«[…]»", "Contacto do pastor", "«[…]»")
    _campo4("Data de ingresso na EBE", "«[__/__/____]»", "Modalidade", "«[Presencial/Híbrida/On-line]»")

    # ====== 2. RESUMO FORMATIVO ======
    _secao_form(doc, "2.  Resumo formativo (síntese por nível)")

    tbl = doc.add_table(rows=1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Nível formativo", "Verbo-chave", "Horas previstas", "Horas cumpridas", "% conclusão", "Situação"]
    widths  = [5.5, 3.5, 3.5, 3.5, 3.0, 6.0]
    hdr = tbl.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].text = ""; hdr[i].width = Cm(widths[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    niveis = [
        ("Nível 1 — Discípulo",     "Conhecer",    "≈ 380 h", "«[]»", "«[]»", "«[Em curso / Concluído / Pendente]»"),
        ("Nível 2 — Crescimento",   "Ser",         "≈ 385 h", "«[]»", "«[]»", "«[]»"),
        ("Nível 3 — Servir",        "Ministério",  "≈ 830 h", "«[]»", "«[]»", "«[]»"),
        ("Nível 4 — Multiplicação", "Reino",       "≈ 700 h", "«[]»", "«[]»", "«[]»"),
    ]
    for r_data in niveis:
        row = tbl.add_row().cells
        for i, v in enumerate(r_data):
            row[i].text = v; row[i].width = Cm(widths[i])
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = FONTE_CORPO; run.font.size = Pt(9)
    # linha total
    row = tbl.add_row().cells
    valores_total = ["TOTAL GERAL", "—", "2.200–2.400 h", "«[]»", "«[]»", "«[]»"]
    for i, v in enumerate(valores_total):
        row[i].text = v; row[i].width = Cm(widths[i])
        _shade_cell(row[i], HEX_SECUNDARIA)
        for p in row[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True; run.font.name = FONTE_TITULO; run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    page_break(doc)

    # ====== 3. REGISTO ACADÊMICO DETALHADO ======
    _secao_form(doc, "3.  Registo acadêmico detalhado")
    paragrafo(doc,
        "Registo nominal e cronológico dos cursos cumpridos pelo(a) aluno(a) "
        "na Escola Bíblica Epignósis, com indicação do instituto, escola, "
        "curso, carga horária, frequência, aproveitamento e data de conclusão.")
    paragrafo(doc,
        "Legenda: A — Aprovado(a)  ·  AC — Aprovado(a) por Compensação  ·  "
        "R — Reprovado(a)  ·  EC — Em Curso  ·  DI — Dispensado(a) por Equivalência.",
        italic=True)

    # === NÍVEL 1 ===
    h2(doc, "Nível 1 — Discípulo (Conhecer)")
    _tabela_registo(doc,
        headers=["Instituto", "Escola", "Curso", "C.H.", "Freq. %", "Nota", "Sit.", "Data conclusão", "Docente"],
        widths=[3.5, 3.5, 5.5, 1.6, 2.0, 1.6, 1.5, 3.0, 4.0],
        n_linhas_vazias=10)

    # === NÍVEL 2 ===
    h2(doc, "Nível 2 — Crescimento (Ser)")
    _tabela_registo(doc,
        headers=["Instituto", "Escola", "Curso", "C.H.", "Freq. %", "Nota", "Sit.", "Data conclusão", "Docente"],
        widths=[3.5, 3.5, 5.5, 1.6, 2.0, 1.6, 1.5, 3.0, 4.0],
        n_linhas_vazias=10)

    page_break(doc)

    # === NÍVEL 3 ===
    h2(doc, "Nível 3 — Servir (Ministério)")
    _tabela_registo(doc,
        headers=["Instituto", "Escola", "Curso", "C.H.", "Freq. %", "Nota", "Sit.", "Data conclusão", "Docente"],
        widths=[3.5, 3.5, 5.5, 1.6, 2.0, 1.6, 1.5, 3.0, 4.0],
        n_linhas_vazias=12)

    # === NÍVEL 4 ===
    h2(doc, "Nível 4 — Multiplicação (Reino)")
    _tabela_registo(doc,
        headers=["Instituto", "Escola", "Curso", "C.H.", "Freq. %", "Nota", "Sit.", "Data conclusão", "Docente"],
        widths=[3.5, 3.5, 5.5, 1.6, 2.0, 1.6, 1.5, 3.0, 4.0],
        n_linhas_vazias=10)

    page_break(doc)

    # ====== 4. EQUIVALÊNCIAS ======
    _secao_form(doc, "4.  Equivalências e validações")
    paragrafo(doc,
        "Cursos ou formações anteriormente realizados pelo(a) aluno(a) em "
        "outras instituições, e aproveitados como equivalência académica "
        "pela Coordenação Acadêmica da Escola Bíblica Epignósis.")
    _tabela_registo(doc,
        headers=["Instituição de origem", "Curso original", "Ano", "C.H.", "Equivalência atribuída na EBE", "Acto de aprovação"],
        widths=[5.5, 5.5, 2.0, 1.8, 7.0, 4.0],
        n_linhas_vazias=6)

    # ====== 5. ESTÁGIOS / PRÁTICAS ======
    _secao_form(doc, "5.  Estágios e práticas ministeriais")
    paragrafo(doc,
        "Registo de actividades práticas, estágios supervisionados, "
        "projectos ministeriais e missões realizadas no decorrer da formação.")
    _tabela_registo(doc,
        headers=["Tipo de actividade", "Local", "Supervisor", "C.H.", "Período", "Avaliação"],
        widths=[5.5, 5.5, 4.5, 1.8, 4.0, 4.5],
        n_linhas_vazias=6)

    page_break(doc)

    # ====== 6. CERTIFICAÇÕES EMITIDAS ======
    _secao_form(doc, "6.  Certificações emitidas")
    paragrafo(doc,
        "Relação dos certificados e diplomas emitidos pela Escola Bíblica "
        "Epignósis em favor do(a) aluno(a):")
    _tabela_registo(doc,
        headers=["N.º registo", "Tipo de certificado", "Objecto (apostila / curso / escola / instituto / programa)", "Data emissão", "Assinaturas"],
        widths=[3.5, 4.5, 11.0, 3.5, 4.0],
        n_linhas_vazias=6)

    # ====== 7. OBSERVAÇÕES PEDAGÓGICAS / PASTORAIS ======
    _secao_form(doc, "7.  Observações pedagógicas e pastorais")
    paragrafo(doc, "Notas relevantes sobre o desempenho académico, espiritual e ministerial do(a) aluno(a):")
    for _ in range(8):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)

    # ====== 8. SITUAÇÃO ACADÊMICA ACTUAL ======
    _secao_form(doc, "8.  Situação acadêmica actual")
    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _campo4(rot1, val1, rot2, val2):
        row = tbl.add_row().cells
        row[0].text = rot1; row[1].text = val1
        row[2].text = rot2; row[3].text = val2
        _shade_cell(row[0], "E8F1EC"); _shade_cell(row[2], "E8F1EC")
        for idx in (0, 2):
            for p in row[idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(9)
                    r.font.color.rgb = COR_PRIMARIA
        for idx in (1, 3):
            for p in row[idx].paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(9)

    _campo4("Situação acadêmica", "«[Activo · Trancado · Concluído · Desligado]»", "Média geral acumulada", "«[…]»")
    _campo4("Frequência média", "«[…] %»", "Horas totais cumpridas", "«[…]»")
    _campo4("Trilha actual", "«[…]»", "Curso em andamento", "«[…]»")
    _campo4("Mentor designado", "«[…]»", "Coordenador de referência", "«[…]»")

    page_break(doc)

    # ====== 9. ATESTAÇÃO INSTITUCIONAL ======
    _secao_form(doc, "9.  Atestação institucional")

    paragrafo(doc,
        "A Escola Bíblica Epignósis, através da sua Secretaria Acadêmica e "
        "da sua Coordenação Acadêmica, atesta que o presente Histórico "
        "Escolar reflecte fielmente o percurso académico do(a) aluno(a) "
        "acima identificado(a), conforme registos oficiais existentes nesta "
        "instituição, em conformidade com o Regimento Acadêmico (EBE-DOC-004) "
        "e o Projecto Pedagógico Oficial (EBE-DOC-003).")

    citacao(doc,
        "Procura apresentar-te a Deus aprovado, como obreiro que não tem de "
        "que se envergonhar, que maneja bem a palavra da verdade.",
        "2 Timóteo 2.15")

    doc.add_paragraph()
    doc.add_paragraph()

    # Assinaturas — 3 colunas
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in tbl.rows[0].cells:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("__________________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    rotulos = [
        ("«[Local], «[__]» de «[mês]» de «[ano]»", "Data de emissão"),
        ("«[Nome do(a) Secretário(a) Acadêmico(a)]»", "Secretaria Acadêmica"),
        ("«[Nome do(a) Coordenador(a) Acadêmico(a)]»", "Coordenação Acadêmica"),
    ]
    for i, (nome, cargo) in enumerate(rotulos):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(nome)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(cargo)
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
        r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    doc.add_paragraph()
    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CARIMBO INSTITUCIONAL")
    r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA

    # Caixa para carimbo
    tbl_c = doc.add_table(rows=1, cols=1)
    tbl_c.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_c.rows[0].cells[0].width = Cm(8)
    cell = tbl_c.rows[0].cells[0]
    for _ in range(6):
        cell.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de validade institucional. Cópia oficial verificável junto à Secretaria Acadêmica\n"
                  "mediante apresentação do n.º de registo: EBE-HIS-«[ano]»-«[n.º]».")
    r.font.name = FONTE_CORPO; r.font.size = Pt(8); r.font.italic = True
    r.font.color.rgb = COR_CITACAO

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conhecer a Deus.  Viver a Palavra.  Manifestar o Reino.  ·  Soli Deo Gloria.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(9); r.font.italic = True
    r.font.color.rgb = COR_PRIMARIA

    out = os.path.join(os.path.dirname(__file__), "EBE-HIS-001_Historico_Escolar.docx")
    doc.save(out)
    print("OK:", out)


if __name__ == "__main__":
    gerar()
