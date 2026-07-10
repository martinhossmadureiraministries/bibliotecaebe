"""
COMPÊNDIO INSTITUCIONAL — reúne os 8 documentos num único .docx.
Estratégia: reusar as funções de geração de cada doc, mas alimentando
o MESMO objecto Document, com capa institucional unificada na frente.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def capa_compendio(doc):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMPÊNDIO INSTITUCIONAL")
    r.font.name = FONTE_TITULO; r.font.size = Pt(13)
    r.font.color.rgb = COR_SECUNDARIA; r.font.bold = True

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ἐπίγνωσις")
    r.font.name = FONTE_TITULO; r.font.size = Pt(22)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color="2E7D4F", size=8)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conhecer a Deus.\nViver a Palavra.\nManifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(15)
    r.font.italic = True; r.font.color.rgb = COR_PRIMARIA

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REUNIÃO OFICIAL DOS DOCUMENTOS INSTITUCIONAIS\n"
                  "(EBE-DOC-001 a EBE-DOC-008)")
    r.font.name = FONTE_CORPO; r.font.size = Pt(11)
    r.font.color.rgb = COR_TEXTO

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(); _add_horizontal_line(p, color="2E7D4F", size=4)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Edição oficial · 2026")
    r.font.name = FONTE_CORPO; r.font.size = Pt(10); r.font.color.rgb = COR_CITACAO

    page_break(doc)


def sumario_geral(doc):
    h1(doc, "Sumário Geral do Compêndio")
    paragrafo(doc, "Este compêndio reúne, em um único volume, os oito documentos institucionais oficiais da Escola Bíblica Epignósis:")
    docs = [
        ("EBE-DOC-001", "Identidade Institucional — Missão, Visão, Valores e Filosofia de Ensino"),
        ("EBE-DOC-002", "Declaração de Fé Institucional"),
        ("EBE-DOC-003", "Projecto Pedagógico Oficial (PPO)"),
        ("EBE-DOC-004", "Regimento Acadêmico"),
        ("EBE-DOC-005", "Arquitectura Oficial"),
        ("EBE-DOC-006", "Mapa Oficial de Cursos"),
        ("EBE-DOC-007", "Duração Oficial dos Cursos"),
        ("EBE-DOC-008", "Sistema de Pré-Requisitos"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(["Código", "Documento"]):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _shade_cell(hdr[i], "1B3A5C")
    for cod, t in docs:
        row = tbl.add_row().cells
        row[0].text = cod; row[1].text = t
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = FONTE_CORPO; run.font.size = Pt(11)

    paragrafo(doc, "")
    paragrafo(doc,
        "Cada documento mantém numeração própria e pode ser consultado de "
        "forma autónoma. Os documentos são complementares e devem ser lidos "
        "em conjunto para uma visão integral da identidade institucional, "
        "do currículo, da governança e da filosofia educativa da Escola "
        "Bíblica Epignósis.")
    page_break(doc)


def divisor(doc, codigo, titulo):
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(codigo)
    r.font.name = FONTE_TITULO; r.font.size = Pt(13)
    r.font.color.rgb = COR_SECUNDARIA; r.font.bold = True

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo.upper())
    r.font.name = FONTE_TITULO; r.font.size = Pt(24); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color="2E7D4F", size=4)
    page_break(doc)


# === Para reusar, importamos cada gerador, mas adaptamos: os geradores
# salvam um doc separado. Em vez disso, recriamos um "corpo" em funções
# alternativas dentro de cada módulo — solução pragmática: importamos as
# funções de conteúdo e reescrevemos aqui um modo "compendiado".
# Mais simples: instalamos um helper que executa os 8 scripts e em seguida
# faz merge usando docxcompose.

if __name__ == "__main__":
    print("Use compendio_merge.py para gerar o compêndio.")
