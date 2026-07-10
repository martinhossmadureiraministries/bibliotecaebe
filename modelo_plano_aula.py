"""
MODELO OFICIAL DE PLANO DE AULA — Escola Bíblica Epignósis
Formulário .docx editável para uso do docente.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line


def _campo(tbl, rotulo, valor=""):
    row = tbl.add_row().cells
    row[0].text = rotulo; row[1].text = valor
    _shade_cell(row[0], "E8F1EC")
    for p in row[0].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
            r.font.color.rgb = COR_PRIMARIA
    for p in row[1].paragraphs:
        for r in p.runs:
            r.font.name = FONTE_CORPO; r.font.size = Pt(10)


def _secao_form(doc, titulo_secao):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"  {titulo_secao.upper()}")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEX_PRIMARIA)
    pPr.append(shd)


def _linhas(doc, n=3):
    for _ in range(n):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)


def gerar():
    doc = novo_documento("Plano de Aula", "EBE-FRM-PLN")

    # CAPA
    doc.add_paragraph()
    inserir_logo(doc, LOGO_PATH, largura_cm=5.0)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ἐπίγνωσις · Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(9)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("PLANO DE AULA")
    r.font.name = FONTE_TITULO; r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Instrumento pedagógico do(a) docente")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    # 1. IDENTIFICAÇÃO
    _secao_form(doc, "1.  Identificação da aula")
    tbl = doc.add_table(rows=0, cols=2)
    _campo(tbl, "Docente responsável", "«[…]»")
    _campo(tbl, "Instituto", "«[…]»")
    _campo(tbl, "Escola", "«[…]»")
    _campo(tbl, "Curso", "«[…]»")
    _campo(tbl, "Módulo", "«[N.º — Nome]»")
    _campo(tbl, "Apostila / Aula", "«[N.º — Título]»")
    _campo(tbl, "Carga horária prevista", "«[…] minutos»")
    _campo(tbl, "Data prevista", "«[__/__/____]»")
    _campo(tbl, "Turma / Modalidade", "«[…]»  ·  «[Presencial / Híbrida / On-line]»")
    _campo(tbl, "Local / Plataforma", "«[…]»")
    _campo(tbl, "Nível formativo", "«[Discípulo / Crescimento / Servir / Multiplicação]»")
    _campo(tbl, "Pré-requisitos exigidos", "«[…]»")

    # 2. TEMA E TEXTO
    _secao_form(doc, "2.  Tema central e texto bíblico")
    tbl = doc.add_table(rows=0, cols=2)
    _campo(tbl, "Tema central da aula", "«[Um conceito central, em uma frase]»")
    _campo(tbl, "Texto-base (ARC)", "«[Livro Cap.versículos]»")
    _campo(tbl, "Versículo-chave (ARC)", "«[Livro Cap.v]»")
    _campo(tbl, "Palavras-chave teológicas", "«[3 a 5 termos]»")

    # 3. OBJECTIVOS
    _secao_form(doc, "3.  Objectivos de aprendizagem (4 eixos)")
    paragrafo(doc, "Ao final desta aula, o(a) aluno(a) será capaz de:")
    h3(doc, "Conhecer (verdade a compreender)")
    _linhas(doc, 2)
    h3(doc, "Crer (convicção a interiorizar)")
    _linhas(doc, 2)
    h3(doc, "Viver (aplicação prática)")
    _linhas(doc, 2)
    h3(doc, "Servir (acção ministerial decorrente)")
    _linhas(doc, 2)

    # 4. CONTEÚDO
    _secao_form(doc, "4.  Conteúdo programático")
    paragrafo(doc, "Tópicos a serem abordados (na ordem prevista):")
    for i in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        r = p.add_run(f"{i+1}.  ")
        r.font.bold = True; r.font.color.rgb = COR_SECUNDARIA
        r2 = p.add_run("«[…]»")
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(11)

    page_break(doc)

    # 5. ROTEIRO DA AULA (com cronograma)
    _secao_form(doc, "5.  Roteiro da aula (cronograma)")
    paragrafo(doc, "Distribuição do tempo, segundo as cinco fases de uma aula Epignósis:")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, t in enumerate(["Fase", "Tempo", "Actividade", "Recursos"]):
        hdr[i].text = ""
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    fases = [
        ("Acolhimento e Oração", "«[5–10 min]»",
         "Saudação, oração inicial, recapitulação da aula anterior, pergunta motivadora.",
         "Bíblia, lista de presenças."),
        ("Apresentação do Tema", "«[10–15 min]»",
         "Introdução do conceito central, contextualização bíblica e prática, exposição do versículo-chave.",
         "Apostila, quadro/projector."),
        ("Desenvolvimento Doutrinário", "«[25–40 min]»",
         "Exposição articulada do conteúdo, leitura comentada de textos bíblicos, definições, esclarecimento de dúvidas comuns.",
         "Apostila, Bíblia (ARC), recursos visuais."),
        ("Aplicação e Diálogo", "«[15–20 min]»",
         "Discussão, exemplos práticos, perguntas dirigidas, partilha de testemunhos, oração de aplicação.",
         "Perguntas norteadoras, exercícios."),
        ("Síntese e Envio", "«[5–10 min]»",
         "Recapitulação, oração final, indicação da leitura para a próxima aula, tarefa de fixação.",
         "Folheto de tarefas, apostila."),
    ]
    for fase in fases:
        row = tbl.add_row().cells
        for i, v in enumerate(fase):
            row[i].text = v
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    # 6. METODOLOGIA E RECURSOS
    _secao_form(doc, "6.  Metodologia e recursos didácticos")
    paragrafo(doc, "Estratégias pedagógicas a utilizar (marque as que se aplicam):")
    estrategias = [
        "Aula expositiva", "Estudo bíblico dirigido", "Leitura comentada",
        "Diálogo socrático", "Debate respeitoso", "Estudo de caso ministerial",
        "Trabalho em pequenos grupos", "Apresentação visual (slides/quadro)",
        "Vídeo de apoio", "Testemunho ou ilustração", "Exercício prático",
        "Momento de oração",
    ]
    for e in estrategias:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f"☐  {e}")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    paragrafo(doc, "Recursos materiais necessários:")
    _linhas(doc, 3)

    # 7. AVALIAÇÃO DA APRENDIZAGEM
    _secao_form(doc, "7.  Avaliação da aprendizagem")
    paragrafo(doc, "Instrumentos previstos nesta aula (marque os que se aplicam):")
    instrumentos = [
        "Participação oral", "Pergunta dirigida", "Questionário escrito",
        "Pequeno trabalho", "Apresentação", "Estudo bíblico para casa",
        "Auto-avaliação", "Projecto ministerial",
    ]
    for i in instrumentos:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f"☐  {i}")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    paragrafo(doc, "Critérios de avaliação:")
    _linhas(doc, 4)

    # 8. INTEGRAÇÃO DEVOCIONAL
    _secao_form(doc, "8.  Integração espiritual e devocional")
    tbl = doc.add_table(rows=0, cols=2)
    _campo(tbl, "Motivo de oração inicial", "«[…]»")
    _campo(tbl, "Versículo para memorização", "«[…]»")
    _campo(tbl, "Aplicação devocional sugerida", "«[…]»")
    _campo(tbl, "Hino / Cântico (opcional)", "«[…]»")
    _campo(tbl, "Encaminhamento pastoral (se houver)", "«[…]»")

    # 9. TAREFAS E PRÓXIMA AULA
    _secao_form(doc, "9.  Tarefas e preparação para a próxima aula")
    paragrafo(doc, "Tarefa de fixação para casa:")
    _linhas(doc, 3)
    paragrafo(doc, "Leitura prévia para a próxima aula:")
    _linhas(doc, 2)
    paragrafo(doc, "Perguntas orientadoras para reflexão:")
    _linhas(doc, 3)

    page_break(doc)

    # 10. RELATÓRIO PÓS-AULA
    _secao_form(doc, "10.  Relatório pós-aula (preenchimento após a aula)")

    tbl = doc.add_table(rows=0, cols=2)
    _campo(tbl, "Data efectiva da aula", "«[__/__/____]»")
    _campo(tbl, "Horário (início/fim)", "«[__:__ / __:__]»")
    _campo(tbl, "N.º de alunos presentes", "«[…]»")
    _campo(tbl, "N.º de alunos previstos", "«[…]»")
    _campo(tbl, "% de frequência", "«[…] %»")

    paragrafo(doc, "Conteúdo efectivamente trabalhado:")
    _linhas(doc, 3)
    paragrafo(doc, "Conteúdo que ficou pendente para a próxima aula:")
    _linhas(doc, 2)
    paragrafo(doc, "Observações sobre a participação e o aproveitamento dos alunos:")
    _linhas(doc, 3)
    paragrafo(doc, "Dificuldades pedagógicas observadas:")
    _linhas(doc, 2)
    paragrafo(doc, "Sugestões de melhoria para a próxima aula:")
    _linhas(doc, 2)
    paragrafo(doc, "Necessidades pastorais identificadas (a encaminhar à Coordenação Espiritual):")
    _linhas(doc, 2)

    citacao(doc,
        "E o que de mim, entre muitas testemunhas, ouviste, confia-o a homens "
        "fiéis, que sejam idóneos para também ensinarem os outros.",
        "2 Timóteo 2.2")

    # ASSINATURAS
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in tbl.rows[0].cells:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("__________________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)
    rot = [
        ("«[Nome do(a) Docente]»", "Docente responsável"),
        ("«[Nome do(a) Coordenador(a) Pedagógico(a)]»", "Coordenação Pedagógica · Visto"),
    ]
    for i, (n, l) in enumerate(rot):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(n)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(l)
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
        r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    selo_final(doc)

    out = os.path.join(os.path.dirname(__file__), "EBE-FRM-002_Plano_de_Aula.docx")
    doc.save(out); print("OK:", out)


if __name__ == "__main__":
    gerar()
