"""
EBE-FRM-005 — Lista de Presença / Chamada Semanal
Escola Bíblica Epignósis
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line


def paisagem(doc):
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def cell_text(cell, text, bold=False, color=None, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FONTE_CORPO
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def header_row(row, labels):
    for cell, label in zip(row.cells, labels):
        _shade_cell(cell, HEX_PRIMARIA)
        cell_text(cell, label, bold=True, color=RGBColor(255,255,255), size=8)


def gerar():
    doc = novo_documento("Lista de Presença", "EBE-FRM-005")
    paisagem(doc)

    add_capa(doc,
        supratitulo="Formulário Académico N.º 5",
        titulo="Lista de Presença / Chamada Semanal",
        subtitulo="Controlo oficial de frequência das turmas da Escola Bíblica Epignósis",
        codigo="EBE-FRM-005", ano="2026")
    add_marco_filosofico(doc)

    h1(doc, "Identificação da Turma", numero=1)
    tbl = doc.add_table(rows=5, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    campos = [
        ("Instituto", "«[Nome do Instituto]»", "Escola", "«[Nome da Escola]»"),
        ("Curso", "«[Nome do Curso]»", "Módulo", "«[Nome/N.º do Módulo]»"),
        ("Apostila / Tema", "«[Tema da aula]»", "Nível", "«[Discípulo/Crescimento/Servir/Multiplicação]»"),
        ("Docente", "«[Nome do(a) docente]»", "Turma", "«[Código/Nome da turma]»"),
        ("Data", "«[__/__/____]»", "Horário", "«[hh:mm – hh:mm]»"),
    ]
    for r, linha in enumerate(campos):
        for c, txt in enumerate(linha):
            cell = tbl.rows[r].cells[c]
            if c in (0,2):
                _shade_cell(cell, "E8F1EC")
                cell_text(cell, txt, bold=True, color=COR_PRIMARIA, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                cell_text(cell, txt, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    h1(doc, "Chamada da Aula", numero=2)
    paragrafo(doc, "Assinale a presença de cada aluno(a) no início da aula e confirme a assinatura ao final. A frequência mínima institucional para certificação é de 75 % das actividades previstas.")
    citacao(doc, "Tudo quanto fizerdes, fazei-o de todo o coração, como ao Senhor, e não aos homens.", "Colossenses 3.23")

    labels = ["N.º", "Nome completo do(a) aluno(a)", "N.º de matrícula", "Contacto", "Presença", "Entrada", "Assinatura", "Observações"]
    tbl = doc.add_table(rows=1, cols=len(labels))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], labels)
    widths = [1.0, 6.2, 2.8, 3.2, 1.8, 1.8, 4.0, 5.2]
    for i in range(1, 31):
        row = tbl.add_row()
        vals = [str(i), "«[Nome do(a) aluno(a)]»" if i == 1 else "", "", "", "□ P  □ F  □ J", "", "", ""]
        for c, val in zip(row.cells, vals):
            cell_text(c, val, size=8, align=WD_ALIGN_PARAGRAPH.CENTER if vals.index(val) in [0,4] else WD_ALIGN_PARAGRAPH.LEFT)
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths[i])

    h1(doc, "Legenda e Critérios de Frequência", numero=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Símbolo", "Significado", "Conta para presença?", "Observação"])
    for data in [
        ("P", "Presente", "Sim", "Aluno(a) participou da aula/actividade."),
        ("F", "Falta", "Não", "Ausência sem justificação aceite."),
        ("J", "Falta justificada", "Conforme decisão", "Deve haver comprovativo ou comunicação formal."),
        ("R", "Reposição", "Sim, se aprovada", "Actividade de recuperação autorizada pela Coordenação."),
    ]:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0,2) else WD_ALIGN_PARAGRAPH.LEFT)

    h1(doc, "Resumo Mensal / Modular de Frequência", numero=4)
    paragrafo(doc, "Use esta grelha quando o módulo ou curso tiver várias aulas. Preencha P, F, J ou R e calcule a percentagem final de frequência.")
    labels = ["N.º", "Nome completo", "Aula 1", "Aula 2", "Aula 3", "Aula 4", "Aula 5", "Aula 6", "Aula 7", "Aula 8", "Total P/R", "Freq. %", "Situação"]
    tbl = doc.add_table(rows=1, cols=len(labels))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], labels)
    for i in range(1, 21):
        row = tbl.add_row().cells
        vals = [str(i), ""] + [""]*8 + ["", "", ""]
        for idx, val in enumerate(vals):
            cell_text(row[idx], val, size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT)

    h1(doc, "Validação", numero=5)
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, label in zip(tbl.rows[0].cells, ["Docente", "Secretaria Acadêmica", "Coordenação Acadêmica"]):
        _shade_cell(cell, "E8F1EC")
        cell_text(cell, label, bold=True, color=COR_PRIMARIA, size=9)
    for cell in tbl.rows[1].cells:
        cell_text(cell, "\n\n__________________________________\nAssinatura e data", size=9)

    selo_final(doc)
    out = os.path.join(os.path.dirname(__file__), "EBE-FRM-005_Lista_de_Presenca.docx")
    doc.save(out)
    print("OK:", out)

if __name__ == "__main__":
    gerar()
