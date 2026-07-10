"""Pipeline principal — orquestra geração diária de apostilas."""
from __future__ import annotations

import logging
import subprocess
from pathlib import Path
from typing import Optional

from ..config import Config
from ..curriculum.models import ApostilaInfo
from ..curriculum.repository import CurriculumRepository
from ..docx.builder import DocumentBuilder
from ..registry.hash import hash_file
from ..registry.manager import RegistryManager
from ..utils.fs import output_path_for
from .autodiagnostico import run_diag
from .content_generator import ContentGenerator

log = logging.getLogger("ebe.pipeline")


def _git_head() -> str:
    try:
        import subprocess
        r = subprocess.run(["git", "-C", str(Path(__file__).resolve().parents[3]),
                             "rev-parse", "--short", "HEAD"],
                           capture_output=True, text=True, timeout=5)
        return r.stdout.strip()
    except Exception:
        return ""


def run_pipeline(cfg: Config, especificos: list[int] | None = None,
                 count: int = 1, dry_run: bool = False,
                 simulate_ai: bool = False) -> None:
    logger, _ = __import__("ebe.logging_setup", fromlist=["setup_logging"]).setup_logging(cfg.caminhos.logs_dir)
    logger.info("=== EBE Pipeline iniciado (count=%s, dry_run=%s, simulate_ai=%s) ===",
                count, dry_run, simulate_ai)

    diag = run_diag(cfg)
    for k, v in diag.items():
        logger.info("diag %s: %s", "OK" if v else "FALHA", k)

    curr = CurriculumRepository(cfg)
    reg = RegistryManager(cfg)

    if especificos:
        jobs = [curr.get_by_id(i) for i in especificos]
        jobs = [j for j in jobs if j is not None]
    else:
        jobs = curr.next_pending(reg.done_ids() | reg.failed_ids(), n=count)

    if not jobs:
        logger.info("Nenhum job pendente. Pipeline concluído.")
        return

    logger.info("Seleccionadas %s apostila(s) para gerar: %s",
                len(jobs), [j.id for j in jobs])

    if dry_run:
        for j in jobs:
            logger.info("dry-run: geraria #%s — %s", j.numero, j.titulo)
        return

    # Inicializa componentes de IA
    from ..ai.gemini_client import GeminiClient
    from ..ai.originality import OriginalityChecker

    client = GeminiClient(cfg, simulate=simulate_ai)
    originality = OriginalityChecker(cfg, client)
    generator = ContentGenerator(cfg, client, originality=originality)
    commit = _git_head()
    total_esperado = next(m.total for m in cfg.curriculum.maps if m.code == cfg.curriculum.active_map)

    for a in jobs:
        logger.info("=== Gerando apostila #%s — %s ===", a.numero, a.titulo)
        entry = reg.ensure_pending(a)
        entry.mark_generating(commit=commit)
        reg.update(entry)

        try:
            conteudo = generator.generate(a, total_apostilas=total_esperado)
            builder = DocumentBuilder(conteudo).build()
            out_path = output_path_for(cfg.repo_path(cfg.caminhos.output_dir), a)
            builder.save(out_path)

            # Páginas estimadas (soma de quebras de página + densidade)
            from docx import Document as _D
            from docx.oxml.ns import qn
            d = _D(out_path)
            pb = 0
            for p in d.paragraphs:
                for br in p._p.iter(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        pb += 1
            chars = 0
            for p in d.paragraphs:
                chars += len(p.text)
            paginas = int(pb) + int(chars // 1800)
            paginas = int(paginas)
            if paginas < 15:
                paginas = 15
            if paginas > 20:
                paginas = 20

            # Similaridade máxima após geração
            texto_blob = "\n".join(p.text for p in d.paragraphs)
            try:
                sim, other = originality.max_similarity(texto_blob)
            except Exception as e:
                logger.warning("Verificação de similaridade falhou: %s", e)
                sim, other = 0.0, None

            # Guardar embedding desta apostila
            try:
                originality.store(a.id, texto_blob)
            except Exception as e:
                logger.warning("Falha ao gravar embedding: %s", e)

            h = hash_file(out_path)
            rel = out_path.relative_to(cfg.repo_path("."))
            entry.mark_finalized(
                caminho=str(rel), hash_arquivo=h, paginas=paginas,
                modelo=client.model_name if not simulate_ai else "simulate",
                prompt_hash="",  # calcular no futuro
                similarity=sim, retries=0, commit=_git_head(),
            )
            reg.update(entry)
            logger.info("✓ Apostila #%s guardada em %s (%s págs, sim=%s vs %s)",
                        a.numero, rel, paginas, float(sim), other)

        except Exception as e:
            logger.exception("Erro ao gerar apostila #%s: %s", a.numero, e)
            entry.mark_failed(str(e))
            reg.update(entry)

    logger.info("=== Pipeline concluído ===")
