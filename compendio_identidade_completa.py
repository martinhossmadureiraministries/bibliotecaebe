"""
EBE-COMP-ID — Compêndio da Identidade Completa da Escola Bíblica Epignósis
Documento-mãe da identidade institucional, doutrinária, pedagógica, organizacional e visual.
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line


def cell_text(cell, text, bold=False, color=None, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FONTE_CORPO
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def header_row(row, labels, size=9):
    for cell, label in zip(row.cells, labels):
        _shade_cell(cell, HEX_PRIMARIA)
        cell_text(cell, label, bold=True, color=RGBColor(255,255,255), size=size, align=WD_ALIGN_PARAGRAPH.CENTER)


def quadro(doc, titulo, texto, cor="E8F1EC"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, cor)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(titulo + " — ")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA
    r2 = p.add_run(texto)
    r2.font.name = FONTE_CORPO
    r2.font.size = Pt(11)


def assinatura_linha(doc, labels):
    tbl = doc.add_table(rows=2, cols=len(labels))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl.rows[0].cells:
        cell_text(cell, "\n\n__________________________________", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for cell, label in zip(tbl.rows[1].cells, labels):
        cell_text(cell, label, bold=True, color=COR_PRIMARIA, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def gerar():
    doc = novo_documento("Compêndio da Identidade Completa", "EBE-COMP-ID")

    add_capa(doc,
        supratitulo="Documento-Mãe da Identidade Institucional",
        titulo="Compêndio da Identidade Completa da EBE",
        subtitulo="Identidade espiritual, doutrinária, pedagógica, académica, organizacional e visual da Escola Bíblica Epignósis",
        codigo="EBE-COMP-ID", ano="2026")
    add_marco_filosofico(doc)

    h1(doc, "Nota de Promulgação")
    paragrafo(doc, "O presente Compêndio da Identidade Completa da Escola Bíblica Epignósis reúne, sistematiza e harmoniza os elementos essenciais que definem a instituição perante Deus, perante a Igreja, perante os seus alunos, docentes, colaboradores e parceiros ministeriais.")
    paragrafo(doc, "Este documento não substitui os documentos institucionais específicos já aprovados, mas funciona como síntese superior da identidade EBE, servindo como referência para comunicação, formação interna, alinhamento doutrinário, acolhimento de novos membros, orientação pedagógica e salvaguarda da unidade institucional.")
    citacao(doc, "Até que todos cheguemos à unidade da fé e ao conhecimento do Filho de Deus, a varão perfeito, à medida da estatura completa de Cristo.", "Efésios 4.13")
    quadro(doc, "Princípio de unidade", "A identidade da EBE não é meramente gráfica, administrativa ou académica. Ela é, antes de tudo, uma identidade espiritual: centrada em Cristo, fundada nas Escrituras e orientada para a edificação da Igreja e a manifestação do Reino de Deus.")

    h1(doc, "Sumário Executivo da Identidade EBE")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Elemento", "Formulação institucional"])
    resumo = [
        ("Nome", "Escola Bíblica Epignósis — EBE."),
        ("Lema", "Conhecer a Deus. Viver a Palavra. Manifestar o Reino."),
        ("Marco filosófico", "O verdadeiro conhecimento de Deus transforma a mente, o coração e a vida."),
        ("Palavra-chave", "ἐπίγνωσις — conhecimento pleno, profundo, experiencial e transformador de Deus."),
        ("Versão bíblica de referência", "Almeida Revista e Corrigida — ARC."),
        ("Eixos pedagógicos", "Conhecer · Crer · Viver · Servir."),
        ("Níveis formativos", "Discípulo · Crescimento · Servir · Multiplicação."),
        ("Institutos fundadores", "10 institutos, organizando a formação bíblica, teológica, espiritual, ministerial e missionária."),
        ("Carga formativa completa", "2.200–2.400 horas, aproximadamente."),
        ("Paleta visual", "Azul-marinho #1B3A5C · Verde #2E7D4F · Dourado #C9A14B."),
    ]
    for a,b in resumo:
        row = tbl.add_row().cells
        cell_text(row[0], a, bold=True, color=COR_PRIMARIA, size=9.5)
        cell_text(row[1], b, size=9.5)

    page_break(doc)

    h1(doc, "Capítulo I — Identidade Espiritual e Teológica", numero=1)
    h2(doc, "1.1. Centro cristológico")
    paragrafo(doc, "A Escola Bíblica Epignósis confessa que Jesus Cristo é o centro da revelação bíblica, o Senhor da Igreja, o Salvador do mundo e o fundamento de todo ensino cristão. Nenhum conteúdo, método, curso, apostila, projecto ou actividade institucional pode ser considerado fiel se não conduzir a Cristo e à obediência à Sua Palavra.")
    citacao(doc, "Porque ninguém pode pôr outro fundamento, além do que já está posto, o qual é Jesus Cristo.", "1 Coríntios 3.11")

    h2(doc, "1.2. Autoridade das Escrituras")
    paragrafo(doc, "A EBE reconhece as Sagradas Escrituras como Palavra de Deus, inspirada, suficiente, verdadeira e normativa para a fé, a doutrina, a vida, a missão e a prática ministerial. A Bíblia não é apenas objecto de estudo; é regra de fé, obediência e transformação.")
    citacao(doc, "Toda Escritura divinamente inspirada é proveitosa para ensinar, para redarguir, para corrigir, para instruir em justiça.", "2 Timóteo 3.16")

    h2(doc, "1.3. Dependência do Espírito Santo")
    paragrafo(doc, "A EBE afirma que o verdadeiro conhecimento de Deus não nasce apenas do esforço intelectual, mas da iluminação e acção do Espírito Santo. Estudo sem oração tende à aridez; oração sem estudo tende à superficialidade. A formação Epignósis integra mente renovada, coração rendido e vida obediente.")
    citacao(doc, "Mas aquele Consolador, o Espírito Santo, que o Pai enviará em meu nome, esse vos ensinará todas as coisas.", "João 14.26")

    h2(doc, "1.4. Vocação eclesial")
    paragrafo(doc, "A EBE existe para servir à Igreja de Cristo, e não para a substituir. A Escola honra a igreja local, reconhece a importância da autoridade pastoral e busca formar discípulos e líderes que retornem às suas comunidades mais maduros, bíblicos, humildes e frutíferos.")
    quadro(doc, "Confissão eclesial", "A Escola Bíblica Epignósis é uma escola para a Igreja, com a Igreja e a serviço da Igreja.")

    h1(doc, "Capítulo II — Nome, Etimologia e Vocação Epignósis", numero=2)
    paragrafo(doc, "O nome Epignósis deriva do termo grego ἐπίγνωσις (epígnōsis), frequentemente utilizado no Novo Testamento para indicar conhecimento pleno, profundo, relacional, experiencial e transformador. A EBE não se contenta com informação religiosa; busca formação integral em Cristo.")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Termo", "Sentido", "Implicação para a EBE"])
    termos = [
        ("γνῶσις — gnōsis", "Conhecimento, informação, compreensão.", "Importante, mas insuficiente se não gerar vida."),
        ("ἐπίγνωσις — epígnōsis", "Conhecimento pleno, profundo, experiencial e transformador.", "Modelo de formação desejado: mente, coração e prática transformados."),
    ]
    for data in termos:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=9.5)
    citacao(doc, "Para que os seus corações sejam consolados, e estejam unidos em caridade e enriquecidos da plenitude da inteligência, para conhecimento do mistério de Deus e Pai, e de Cristo.", "Colossenses 2.2")
    citacao(doc, "Para conhecê-lo, e à virtude da sua ressurreição, e à comunicação de suas aflições, sendo feito conforme à sua morte.", "Filipenses 3.10")

    h1(doc, "Capítulo III — Lema, Missão, Visão e Marco Filosófico", numero=3)
    h2(doc, "3.1. Lema institucional")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conhecer a Deus.  Viver a Palavra.  Manifestar o Reino.")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = COR_PRIMARIA

    h2(doc, "3.2. Missão")
    paragrafo(doc, "Conduzir cada aluno ao conhecimento pleno de Deus por meio das Escrituras, da acção do Espírito Santo e da prática do Evangelho, formando discípulos maduros, líderes íntegros, servos frutíferos e multiplicadores do Reino.")

    h2(doc, "3.3. Visão")
    paragrafo(doc, "Ser uma referência na formação bíblica, teológica, espiritual, ministerial e missionária de discípulos de Cristo, servindo à Igreja com excelência, fidelidade doutrinária, profundidade espiritual e impacto transformador.")

    h2(doc, "3.4. Marco filosófico oficial")
    paragrafo(doc, "Acreditamos que o verdadeiro conhecimento de Deus transforma a mente pela verdade das Escrituras, o coração pela acção do Espírito Santo e a vida pelo compromisso de viver e anunciar o Evangelho de Jesus Cristo.", italic=True)

    h1(doc, "Capítulo IV — Valores Institucionais", numero=4)
    valores = [
        ("Fidelidade bíblica", "Toda formação deve nascer das Escrituras e regressar às Escrituras."),
        ("Cristocentrismo", "Cristo é o centro do ensino, da vida espiritual, da missão e da esperança cristã."),
        ("Dependência do Espírito Santo", "Sem o Espírito, o estudo torna-se letra sem vida; com o Espírito, a verdade transforma."),
        ("Santidade prática", "Conhecimento verdadeiro produz obediência, carácter e vida limpa diante de Deus."),
        ("Excelência académica", "O Reino merece estudo sério, boa escrita, método, revisão e rigor."),
        ("Serviço humilde", "Todo dom, conhecimento e título existem para servir, não para exaltar o homem."),
        ("Unidade da Igreja", "A EBE honra a diversidade denominacional dentro da fidelidade aos fundamentos da fé."),
        ("Missão e multiplicação", "Todo aluno é formado para frutificar, discipular e multiplicar o que recebeu."),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Valor", "Descrição"])
    for a,b in valores:
        row = tbl.add_row().cells
        cell_text(row[0], a, bold=True, color=COR_PRIMARIA, size=9.5)
        cell_text(row[1], b, size=9.5)

    h1(doc, "Capítulo V — Declaração Doutrinária Sintética", numero=5)
    paragrafo(doc, "A identidade doutrinária completa da EBE está formalmente definida na Declaração de Fé Institucional (EBE-DOC-002). Este Compêndio apresenta a síntese orientadora dos seus principais eixos confessionais.")
    artigos = [
        ("Escrituras", "Cremos na inspiração divina, autoridade, suficiência e veracidade das Sagradas Escrituras."),
        ("Deus Triúno", "Cremos em um só Deus, eternamente existente em três pessoas: Pai, Filho e Espírito Santo."),
        ("Jesus Cristo", "Cremos na divindade, humanidade, morte expiatória, ressurreição corporal, ascensão e retorno glorioso de Cristo."),
        ("Espírito Santo", "Cremos na pessoa e obra do Espírito Santo, que convence, regenera, santifica, capacita e guia a Igreja."),
        ("Ser humano e pecado", "Cremos que o ser humano foi criado à imagem de Deus, caiu em pecado e necessita de redenção."),
        ("Salvação", "Cremos na salvação pela graça, mediante a fé em Jesus Cristo, com frutos de arrependimento e nova vida."),
        ("Igreja", "Cremos na Igreja como Corpo de Cristo, chamada à adoração, comunhão, discipulado, serviço e missão."),
        ("Baptismo e Ceia", "Cremos nas ordenanças instituídas por Cristo, praticadas com reverência e fidelidade bíblica."),
        ("Vida cristã", "Cremos na santificação, no testemunho, na oração, na vida devocional e na prática do amor."),
        ("Últimas coisas", "Cremos no retorno de Cristo, na ressurreição, no juízo final e na consumação do Reino de Deus."),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Doutrina", "Síntese"])
    for a,b in artigos:
        row = tbl.add_row().cells
        cell_text(row[0], a, bold=True, color=COR_PRIMARIA, size=9)
        cell_text(row[1], b, size=9)
    citacao(doc, "Tem cuidado de ti mesmo e da doutrina. Persevera nestes deveres; porque, fazendo isto, te salvarás, tanto a ti mesmo como aos que te ouvem.", "1 Timóteo 4.16")

    page_break(doc)

    h1(doc, "Capítulo VI — Identidade Pedagógica", numero=6)
    h2(doc, "6.1. O objectivo pedagógico maior")
    paragrafo(doc, "A pedagogia Epignósis procura formar discípulos integrais, capazes de conhecer biblicamente, crer doutrinariamente, viver coerentemente e servir frutiferamente. O processo educativo não visa apenas certificação, mas transformação.")

    h2(doc, "6.2. Os quatro eixos pedagógicos")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Eixo", "Ênfase", "Pergunta formativa", "Peso avaliativo"])
    eixos = [
        ("Conhecer", "Verdade bíblica e compreensão doutrinária.", "O aluno compreendeu fielmente?", "40 %"),
        ("Crer", "Convicção, fé e adesão à verdade.", "O aluno crê e confessa com coerência?", "20 %"),
        ("Viver", "Carácter, santidade e prática quotidiana.", "O aluno vive o que aprende?", "20 %"),
        ("Servir", "Dons, missão e ministério.", "O aluno coloca o aprendido a serviço do Reino?", "20 %"),
    ]
    for data in eixos:
        row = tbl.add_row().cells
        for i, val in enumerate(data):
            cell_text(row[i], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0,3) else WD_ALIGN_PARAGRAPH.LEFT)

    h2(doc, "6.3. Método Epignósis de estudo")
    lista(doc, [
        "ORAR — pedir iluminação ao Espírito Santo.",
        "LER — aproximar-se do texto bíblico com atenção e reverência.",
        "ESTUDAR — investigar contexto, doutrina, termos e aplicações.",
        "ANOTAR — registar descobertas, dúvidas e convicções.",
        "MEMORIZAR — guardar a Palavra no coração.",
        "APLICAR — traduzir a verdade em obediência concreta.",
        "ENSINAR — partilhar fielmente o que foi recebido.",
    ], ordenada=True)
    citacao(doc, "Escondi a tua palavra no meu coração, para eu não pecar contra ti.", "Salmo 119.11")

    h1(doc, "Capítulo VII — Arquitectura Académica Oficial", numero=7)
    paragrafo(doc, "A arquitectura da Escola Bíblica Epignósis organiza o ensino de modo progressivo, modular e certificável, permitindo que cada aluno avance de fundamentos cristãos até à maturidade ministerial e multiplicadora.")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Nível estrutural", "Descrição", "Função", "Exemplo"])
    estrutura = [
        ("1. EBE", "Instituição-mãe", "Guardar identidade, doutrina, padrões e certificação.", "Escola Bíblica Epignósis"),
        ("2. Instituto", "Grande domínio formativo", "Agrupar escolas por área teológica/ministério.", "Instituto de Ciências Bíblicas"),
        ("3. Escola", "Área específica", "Organizar cursos afins.", "Escola de Hermenêutica"),
        ("4. Programa", "Percurso formativo integrado", "Agrupar cursos para finalidade específica.", "Programa de Formação Ministerial"),
        ("5. Curso", "Unidade curricular ampla", "Desenvolver competência bíblica/ministerial.", "Introdução à Hermenêutica"),
        ("6. Módulo", "Bloco temático", "Dividir o curso em temas macro.", "Regras de Interpretação"),
        ("7. Apostila", "Unidade mínima de ensino", "Ensinar um conceito central em 1–3 horas.", "O Contexto Imediato"),
    ]
    for data in estrutura:
        row = tbl.add_row().cells
        for i,val in enumerate(data):
            cell_text(row[i], val, size=8.5)

    h1(doc, "Capítulo VIII — Níveis Formativos", numero=8)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Nível", "Verbo-chave", "Ênfase", "Carga aproximada"])
    niveis = [
        ("Nível 1 — Discípulo", "CONHECER", "Fundamentos da fé cristã, novo nascimento, oração, santidade e discipulado inicial.", "≈ 380 h"),
        ("Nível 2 — Crescimento", "SER", "Maturidade doutrinária, hermenêutica, vida espiritual e formação do carácter.", "≈ 385 h"),
        ("Nível 3 — Servir", "MINISTÉRIO", "Capacitação ministerial, ensino, pregação, liderança, dons e prática de serviço.", "≈ 830 h"),
        ("Nível 4 — Multiplicação", "REINO", "Formação de líderes multiplicadores, missão, pesquisa, excelência e impacto no Reino.", "≈ 700 h"),
    ]
    for data in niveis:
        row = tbl.add_row().cells
        for i,val in enumerate(data):
            cell_text(row[i], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (1,3) else WD_ALIGN_PARAGRAPH.LEFT)
    paragrafo(doc, "A Formação Completa Epignósis compreende aproximadamente 2.200–2.400 horas, conforme trilhas, equivalências, práticas e programas complementares.", italic=True)

    h1(doc, "Capítulo IX — Os Dez Institutos Fundadores", numero=9)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["N.º", "Instituto", "Vocação formativa"])
    institutos = [
        ("1", "Instituto de Formação Cristã", "Fundamentos, discipulado, vida cristã e identidade em Cristo."),
        ("2", "Instituto de Ciências Bíblicas", "Bíblia, panorama, exegese, hermenêutica e interpretação fiel."),
        ("3", "Instituto de Ciências Teológicas", "Doutrinas cristãs, teologia sistemática, bíblica e histórica."),
        ("4", "Instituto de Formação Espiritual", "Vida devocional, carácter, disciplinas espirituais e maturidade."),
        ("5", "Instituto Ministerial", "Pregação, ensino, aconselhamento, administração e prática ministerial."),
        ("6", "Instituto do Reino e Poder", "Reino de Deus, dons espirituais, autoridade, oração e poder para servir."),
        ("7", "Instituto dos Cinco Ministérios", "Apóstolos, profetas, evangelistas, pastores e mestres para edificação da Igreja."),
        ("8", "Instituto de Missões", "Evangelização, missões locais, urbanas, transculturais e plantação de igrejas."),
        ("9", "Instituto de Liderança e Multiplicação", "Liderança servidora, mentoria, gestão ministerial e formação de líderes."),
        ("10", "Instituto de Pesquisa Bíblica e Excelência", "Metodologia, investigação, escrita, produção de materiais e excelência académica."),
    ]
    for data in institutos:
        row = tbl.add_row().cells
        for i,val in enumerate(data):
            cell_text(row[i], val, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)

    page_break(doc)

    h1(doc, "Capítulo X — Identidade Académica e Avaliativa", numero=10)
    paragrafo(doc, "A EBE adopta avaliação integral, respeitando o desenvolvimento cognitivo, doutrinário, espiritual e ministerial do aluno. A nota final não mede apenas informação memorizada, mas maturidade formativa nas quatro dimensões institucionais.")
    h2(doc, "10.1. Escala institucional de notas")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Intervalo", "Menção", "Sentido formativo"])
    escala = [
        ("9,0–10,0", "Excelente", "Domínio elevado, clareza, maturidade, aplicação e serviço com excelência."),
        ("7,5–8,9", "Muito bom", "Bom domínio e aplicação consistente."),
        ("6,0–7,4", "Satisfatório", "Cumpre os requisitos essenciais de aprovação."),
        ("5,0–5,9", "Suficiente", "Aproveitamento parcial, insuficiente para aprovação institucional."),
        ("0,0–4,9", "Insuficiente", "Não atingiu os objectivos mínimos."),
    ]
    for data in escala:
        row = tbl.add_row().cells
        for i,val in enumerate(data):
            cell_text(row[i], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT)
    paragrafo(doc, "Média mínima institucional de aprovação: 6,0. Frequência mínima ordinária: 75 %.", italic=True)

    h2(doc, "10.2. Certificação progressiva")
    lista(doc, [
        "Certificado de Apostila — quando aplicável à unidade mínima de estudo.",
        "Certificado de Módulo — após conclusão de conjunto de apostilas.",
        "Certificado de Curso — após conclusão de módulos e avaliação final.",
        "Certificado de Escola — após conclusão dos cursos de determinada escola.",
        "Certificado de Instituto — após conclusão das escolas de um instituto.",
        "Certificado de Programa — quando o aluno conclui uma trilha/programa integrado.",
        "Diploma de Formação Completa Epignósis — após conclusão integral prevista." ,
    ])

    h1(doc, "Capítulo XI — Perfis Institucionais", numero=11)
    h2(doc, "11.1. Perfil do aluno Epignósis")
    lista(doc, [
        "Coração ensinável e reverente diante das Escrituras.",
        "Disciplina no estudo e na vida devocional.",
        "Compromisso com a igreja local.",
        "Honestidade académica e rejeição ao plágio.",
        "Conduta cristã coerente dentro e fora da Escola.",
        "Disposição para servir e multiplicar o que aprende.",
    ])
    h2(doc, "11.2. Perfil do docente Epignósis")
    lista(doc, [
        "Vocação reconhecida para ensinar.",
        "Adesão formal à Declaração de Fé Institucional.",
        "Preparo bíblico-teológico compatível com a matéria ensinada.",
        "Vida espiritual e testemunho público coerentes.",
        "Capacidade didáctica, pastoralidade e humildade.",
        "Submissão aos processos de revisão pedagógica e doutrinária." ,
    ])
    citacao(doc, "E o que de mim, entre muitas testemunhas, ouviste, confia-o a homens fiéis, que sejam idóneos para também ensinarem os outros.", "2 Timóteo 2.2")

    h1(doc, "Capítulo XII — Identidade Organizacional e Governança", numero=12)
    paragrafo(doc, "A governança da EBE deve preservar a unidade doutrinária, a qualidade pedagógica, a organização académica, a integridade administrativa e o cuidado espiritual dos alunos e docentes.")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Órgão / função", "Responsabilidade essencial"])
    orgaos = [
        ("Direcção Geral", "Guardar a visão, representar institucionalmente a Escola e supervisionar o cumprimento da missão."),
        ("Coordenação Acadêmica", "Gerir calendário, turmas, matrículas, frequência, avaliações, certificações e histórico escolar."),
        ("Coordenação Pedagógica", "Acompanhar metodologia, planos de aula, produção de materiais e formação docente."),
        ("Coordenação Espiritual e Pastoral", "Acompanhar espiritualmente alunos e docentes, especialmente em situações de crise e disciplina restauradora."),
        ("Conselho Doutrinário", "Guardar a fidelidade bíblica e revisar conteúdos, apostilas, cursos e declarações confessionais."),
        ("Secretaria Acadêmica", "Manter registos, documentos, pautas, listas de presença e emissão documental."),
    ]
    for data in orgaos:
        row = tbl.add_row().cells
        cell_text(row[0], data[0], bold=True, color=COR_PRIMARIA, size=9)
        cell_text(row[1], data[1], size=9)

    h1(doc, "Capítulo XIII — Identidade Visual e Comunicação Institucional", numero=13)
    paragrafo(doc, "A identidade visual da Escola Bíblica Epignósis deve comunicar seriedade académica, profundidade bíblica, vida espiritual e serviço ao Reino. A marca não é mero ornamento; é sinal visível de uma vocação espiritual e institucional.")

    h2(doc, "13.1. Logotipo")
    paragrafo(doc, "O logotipo oficial da EBE integra elementos visuais que comunicam a identidade da Escola: livro aberto, cruz, monograma EBE, cores institucionais e a designação Escola Bíblica Epignósis. O seu uso deve respeitar proporção, legibilidade, cores oficiais e contexto de aplicação.")
    if os.path.exists(LOGO_PEQUENO):
        inserir_logo(doc, LOGO_PEQUENO, largura_cm=4.5)

    h2(doc, "13.2. Paleta institucional")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Cor", "HEX", "Função", "Amostra"])
    cores = [
        ("Azul-marinho", "#1B3A5C", "Cor primária: títulos, capas, bordas, autoridade académica.", HEX_PRIMARIA),
        ("Verde institucional", "#2E7D4F", "Cor secundária: subtítulos, linhas, elementos de vida e crescimento.", HEX_SECUNDARIA),
        ("Dourado", "#C9A14B", "Cor terciária: selos, destaques solenes e reconhecimento.", HEX_TERCIARIA),
        ("Verde claro", "#E8F1EC", "Fundo discreto para tabelas e quadros pedagógicos.", "E8F1EC"),
    ]
    for nome,hexv,uso,amostra in cores:
        row = tbl.add_row().cells
        cell_text(row[0], nome, bold=True, color=COR_PRIMARIA, size=9)
        cell_text(row[1], hexv, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[2], uso, size=9)
        _shade_cell(row[3], amostra)
        cell_text(row[3], " ", size=9)

    h2(doc, "13.3. Tipografia")
    paragrafo(doc, "A tipografia oficial adoptada é Garamond, tanto para títulos como para corpo textual, reforçando uma estética académica, clássica, formal e adequada a uma escola bíblica de tradição teológica.")

    h2(doc, "13.4. Tom de comunicação")
    lista(doc, [
        "Académico, mas pastoral.",
        "Firme na doutrina, mas humilde na postura.",
        "Fraterno no acolhimento, mas rigoroso na formação.",
        "Claro, sóbrio, reverente e bíblico.",
        "Preferencialmente em português europeu/Angola, respeitando a realidade linguística local." ,
    ])

    page_break(doc)

    h1(doc, "Capítulo XIV — Sistema Documental Oficial", numero=14)
    paragrafo(doc, "A identidade institucional é preservada também por meio de documentos oficiais, códigos e modelos padronizados. A tabela seguinte resume a base documental da EBE.")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row(tbl.rows[0], ["Código", "Documento", "Função"])
    documentos = [
        ("EBE-DOC-001", "Identidade Institucional", "Define missão, visão, valores, filosofia e identidade base."),
        ("EBE-DOC-002", "Declaração de Fé", "Estabelece a confissão doutrinária oficial."),
        ("EBE-DOC-003", "Projecto Pedagógico Oficial", "Define método, eixos, processos e finalidade formativa."),
        ("EBE-DOC-004", "Regimento Académico", "Normatiza direitos, deveres, avaliação, certificação e disciplina."),
        ("EBE-DOC-005", "Arquitectura Oficial", "Organiza institutos, escolas, cursos, módulos e apostilas."),
        ("EBE-DOC-006", "Mapa de Cursos", "Lista os cursos e escolas da estrutura formativa."),
        ("EBE-DOC-007", "Duração Oficial", "Define carga horária dos cursos e níveis."),
        ("EBE-DOC-008", "Sistema de Pré-Requisitos", "Organiza progressão e dependências formativas."),
        ("EBE-MAN-ALU", "Manual do Aluno", "Guia prático, pastoral e académico do aluno."),
        ("EBE-MAN-DOC", "Manual do Docente", "Guia formativo, pedagógico e pastoral do docente."),
        ("EBE-FRM-001 a 006", "Formulários oficiais", "Matrícula, plano, declaração, avaliação, presença e pauta."),
        ("EBE-MOD-TCC-001", "Modelo de TCC/Monografia", "Modelo universal de trabalho final."),
    ]
    for data in documentos:
        row = tbl.add_row().cells
        for i,val in enumerate(data):
            cell_text(row[i], val, size=8.5)

    h1(doc, "Capítulo XV — Princípios de Uso da Identidade", numero=15)
    h2(doc, "15.1. O que deve ser preservado")
    lista(doc, [
        "O nome Escola Bíblica Epignósis e a sigla EBE devem ser usados de modo digno, formal e coerente.",
        "O lema institucional deve aparecer em documentos oficiais, capas, certificados e materiais de acolhimento.",
        "O logotipo não deve ser distorcido, recortado indevidamente ou usado com baixa legibilidade.",
        "As cores institucionais devem ser respeitadas em documentos, certificados, apresentações e materiais gráficos.",
        "Toda produção doutrinária deve ser coerente com a Declaração de Fé.",
        "Toda produção pedagógica deve respeitar os quatro eixos: Conhecer, Crer, Viver e Servir." ,
    ])
    h2(doc, "15.2. O que deve ser evitado")
    lista(doc, [
        "Usar a marca EBE para iniciativas não aprovadas pela Direcção.",
        "Publicar materiais em nome da EBE sem revisão pedagógica e doutrinária.",
        "Alterar o logotipo, cores, tipografia ou lema sem autorização institucional.",
        "Ensinar doutrinas contrárias à Declaração de Fé Institucional.",
        "Reduzir a EBE a plataforma de promoção pessoal, política ou denominacional." ,
    ])

    h1(doc, "Capítulo XVI — Declaração de Identidade e Compromisso", numero=16)
    paragrafo(doc, "A Escola Bíblica Epignósis declara, diante de Deus e da Igreja de Cristo, que a sua identidade será guardada pela fidelidade às Escrituras, centralidade de Cristo, dependência do Espírito Santo, compromisso com a santidade, excelência académica, serviço humilde e missão multiplicadora.")
    citacao(doc, "Portanto, ide, ensinai todas as nações, baptizando-as em nome do Pai, e do Filho, e do Espírito Santo; ensinando-as a guardar todas as coisas que eu vos tenho mandado.", "Mateus 28.19-20")
    quadro(doc, "Declaração-síntese", "Somos a Escola Bíblica Epignósis: uma comunidade de formação cristã chamada a conhecer Deus profundamente, viver a Palavra fielmente e manifestar o Reino frutiferamente.")

    h2(doc, "Oração institucional")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.right_indent = Cm(1.0)
    r = p.add_run("“Senhor Deus, guarda a Escola Bíblica Epignósis em fidelidade à Tua Palavra. Livra-nos da superficialidade, da vaidade, da confusão doutrinária e da esterilidade espiritual. Faz de nós uma escola que ensina com verdade, ama com sinceridade, vive com santidade e serve com poder. Que cada aluno, docente e colaborador conheça a Ti de modo profundo e manifeste o Teu Reino em todas as áreas da vida. Em nome de Jesus, amém.”")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(12.5)
    r.font.italic = True
    r.font.color.rgb = COR_PRIMARIA

    h2(doc, "Homologação interna")
    paragrafo(doc, "O presente Compêndio poderá ser usado como documento de referência para reuniões de liderança, formação docente, acolhimento de alunos, apresentação pública da Escola e orientação institucional.")
    assinatura_linha(doc, ["Direcção Geral", "Coordenação Acadêmica", "Conselho Doutrinário"])

    selo_final(doc)
    out = os.path.join(os.path.dirname(__file__), "EBE-COMP-ID_Compendio_da_Identidade_Completa.docx")
    doc.save(out)
    print("OK:", out)


if __name__ == "__main__":
    gerar()
