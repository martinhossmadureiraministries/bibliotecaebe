"""
MODELOS OFICIAIS DE CERTIFICADOS — Escola Bíblica Epignósis

Gera um único .docx contendo 7 modelos editáveis em paisagem (A4):

  1. Certificado de Apostila
  2. Certificado de Módulo
  3. Certificado de Curso
  4. Certificado de Escola
  5. Certificado de Instituto
  6. Certificado de Programa de Formação
  7. Diploma de Formação Completa (Diploma Final)
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _ornamento_canto(doc, alinhamento=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph(); p.alignment = alinhamento
    r = p.add_run("✦  ✦  ✦")
    r.font.name = FONTE_TITULO; r.font.size = Pt(12)
    r.font.color.rgb = COR_SECUNDARIA


def _borda_pagina(section):
    """Aplica uma borda decorativa dupla em toda a página (paisagem)."""
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    pgBorders.set(qn("w:display"), "allPages")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "double")
        b.set(qn("w:sz"), "18")
        b.set(qn("w:space"), "24")
        b.set(qn("w:color"), HEX_PRIMARIA)
        pgBorders.append(b)
    # inserir
    sectPr.append(pgBorders)


def _configurar_paisagem(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    # A4 paisagem
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    _borda_pagina(section)


def _capa_certificado(doc, tipo_titulo, corpo_paragrafos, pre_titulo="CERTIFICA-SE QUE", primeira_pagina=False):
    """Renderiza um certificado completo numa página paisagem."""

    # Linha superior decorativa
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(LOGO_PATH, width=Cm(3.5))

    # Cabeçalho institucional
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("ἐπίγνωσις  ·  Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(9)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    # Título do certificado
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(tipo_titulo.upper())
    r.font.name = FONTE_TITULO; r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    # Frase de abertura
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(pre_titulo)
    r.font.name = FONTE_TITULO; r.font.size = Pt(12)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    # Corpo
    for texto, tamanho in corpo_paragrafos:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(texto)
        r.font.name = FONTE_TITULO; r.font.size = Pt(tamanho)
        if tamanho >= 20:
            r.font.bold = True; r.font.color.rgb = COR_PRIMARIA
        elif tamanho >= 14:
            r.font.color.rgb = COR_PRIMARIA
        else:
            r.font.color.rgb = COR_TEXTO

    # Versículo
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("“Conheçamos, e prossigamos em conhecer ao Senhor.”  — Oseias 6.3 (ARC)")
    r.font.name = FONTE_TITULO; r.font.size = Pt(10)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    # Espaço antes das assinaturas
    doc.add_paragraph()

    # Tabela de assinaturas (3 colunas: Local/Data | Diretor | Coordenador)
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    larguras = [Cm(8), Cm(8), Cm(8)]
    for r in tbl.rows:
        for i, cell in enumerate(r.cells):
            cell.width = larguras[i]

    # Linha 1 — espaço para assinatura (linha)
    for i, c in enumerate(tbl.rows[0].cells):
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("__________________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10); r.font.color.rgb = COR_TEXTO

    # Linha 2 — rótulos
    rotulos = [
        ("«[Local]», «[dia]» de «[mês]» de «[ano]»", ""),
        ("«[Nome do Director Geral]»", "Director Geral"),
        ("«[Nome do Coordenador Acadêmico]»", "Coordenação Acadêmica"),
    ]
    for i, (nome, cargo) in enumerate(rotulos):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(nome)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        if cargo:
            p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(cargo)
            r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
            r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    # Rodapé com código
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Registo institucional: EBE-CERT-«[tipo]»-«[ano]»-«[n.º]»   ·   Verificável junto à Secretaria Acadêmica.")
    r.font.name = FONTE_CORPO; r.font.size = Pt(8); r.font.italic = True
    r.font.color.rgb = COR_CITACAO


def _nova_secao_paisagem(doc):
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configurar_paisagem(new_section)
    return new_section


def gerar():
    doc = Document()
    configurar_estilos_base(doc)

    # primeira secção em paisagem
    _configurar_paisagem(doc.sections[0])

    # ===== ÍNDICE DE MODELOS (PORTRAIT) =====
    # Vamos começar pela página índice em PAISAGEM mesmo, sem trocar
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(LOGO_PATH, width=Cm(4.5))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MODELOS OFICIAIS DE CERTIFICAÇÃO")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(26); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    doc.add_paragraph()
    paragrafo(doc, "Este documento reúne os sete modelos oficiais de certificação da Escola Bíblica Epignósis, em ordem progressiva da menor unidade (Apostila) ao Diploma de Formação Completa. Todos os modelos são editáveis: basta substituir os campos entre «[...]» pelos dados próprios do aluno e do curso.")

    tbl = doc.add_table(rows=1, cols=3)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(["N.º", "Modelo", "Quando emitir"]):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _shade_cell(hdr[i], HEX_PRIMARIA)
    rows = [
        ("1", "Certificado de Apostila", "Após a conclusão de uma apostila (1–3 horas)."),
        ("2", "Certificado de Módulo", "Após a conclusão de todas as apostilas de um módulo (4–10 horas)."),
        ("3", "Certificado de Curso", "Após a conclusão de todos os módulos de um curso (20–60 horas)."),
        ("4", "Certificado de Escola", "Após a conclusão de todos os cursos de uma Escola (60–180 horas)."),
        ("5", "Certificado de Instituto", "Após a conclusão de todas as Escolas de um Instituto (200–600 horas)."),
        ("6", "Certificado de Programa de Formação", "Após a conclusão de um nível formativo (Discípulo, Crescimento, Servir ou Multiplicação)."),
        ("7", "Diploma de Formação Completa Epignósis", "Após a conclusão das 2.200–2.400 horas da formação integral."),
    ]
    for r in rows:
        row = tbl.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.name = FONTE_CORPO; run.font.size = Pt(10)

    page_break(doc)

    # ============= 1. APOSTILA =============
    _capa_certificado(doc,
        tipo_titulo="Certificado de Apostila",
        pre_titulo="CERTIFICA-SE QUE",
        corpo_paragrafos=[
            ("«[NOME COMPLETO DO ALUNO]»", 26),
            ("concluiu, com aproveitamento, a apostila", 12),
            ("«[Título da Apostila]»", 18),
            ("Módulo «[N.º — Nome]»  ·  Curso «[Nome]»  ·  Escola «[Nome]»", 11),
            ("Instituto «[Nome]»  ·  Carga horária: «[1–3]» horas.", 11),
        ])

    _nova_secao_paisagem(doc)

    # ============= 2. MÓDULO =============
    _capa_certificado(doc,
        tipo_titulo="Certificado de Módulo",
        pre_titulo="CERTIFICA-SE QUE",
        corpo_paragrafos=[
            ("«[NOME COMPLETO DO ALUNO]»", 26),
            ("concluiu, com aproveitamento, o módulo", 12),
            ("«[Título do Módulo]»", 18),
            ("composto pelas apostilas «[indicação das apostilas]».", 11),
            ("Curso «[Nome]»  ·  Escola «[Nome]»  ·  Instituto «[Nome]».", 11),
            ("Carga horária total: «[4–10]» horas.", 11),
        ])

    _nova_secao_paisagem(doc)

    # ============= 3. CURSO =============
    _capa_certificado(doc,
        tipo_titulo="Certificado de Curso",
        pre_titulo="CERTIFICA-SE QUE",
        corpo_paragrafos=[
            ("«[NOME COMPLETO DO ALUNO]»", 28),
            ("concluiu, com aproveitamento, o curso", 13),
            ("«[Título do Curso]»", 22),
            ("Escola «[Nome]»  ·  Instituto «[Nome]».", 12),
            ("Carga horária total: «[20–60]» horas.", 12),
            ("«[Indicação opcional dos módulos cumpridos.]»", 10),
        ])

    _nova_secao_paisagem(doc)

    # ============= 4. ESCOLA =============
    _capa_certificado(doc,
        tipo_titulo="Certificado de Escola",
        pre_titulo="CERTIFICA-SE QUE",
        corpo_paragrafos=[
            ("«[NOME COMPLETO DO ALUNO]»", 28),
            ("concluiu, com aproveitamento, a", 13),
            ("Escola de «[Nome da Escola]»", 24),
            ("do Instituto «[Nome do Instituto]».", 13),
            ("Carga horária total: «[60–180]» horas.", 12),
            ("Cursos cumpridos: «[lista resumida dos cursos].»", 10),
        ])

    _nova_secao_paisagem(doc)

    # ============= 5. INSTITUTO =============
    _capa_certificado(doc,
        tipo_titulo="Certificado de Instituto",
        pre_titulo="CERTIFICA-SE QUE",
        corpo_paragrafos=[
            ("«[NOME COMPLETO DO ALUNO]»", 28),
            ("concluiu, com aproveitamento, o", 13),
            ("Instituto de «[Nome do Instituto]»", 24),
            ("composto pelas Escolas: «[lista das escolas]».", 11),
            ("Carga horária total: «[200–600]» horas.", 12),
        ])

    _nova_secao_paisagem(doc)

    # ============= 6. PROGRAMA / NÍVEL =============
    _capa_certificado(doc,
        tipo_titulo="Certificado de Programa de Formação",
        pre_titulo="CERTIFICA-SE QUE",
        corpo_paragrafos=[
            ("«[NOME COMPLETO DO ALUNO]»", 28),
            ("concluiu, com aproveitamento, o nível formativo", 12),
            ("«[DISCÍPULO ·  CRESCIMENTO ·  SERVIR ·  MULTIPLICAÇÃO]»", 18),
            ("integrado ao Programa «[nome do programa]».", 12),
            ("Carga horária total acumulada: «[h]» horas.", 12),
            ("Institutos cumpridos: «[lista]».", 10),
        ])

    _nova_secao_paisagem(doc)

    # ============= 7. DIPLOMA FINAL =============
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(LOGO_PATH, width=Cm(4.0))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ἐπίγνωσις  ·  Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(9)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("DIPLOMA")
    r.font.name = FONTE_TITULO; r.font.size = Pt(34); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DE FORMAÇÃO COMPLETA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Confere a")
    r.font.name = FONTE_TITULO; r.font.size = Pt(12)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("«[NOME COMPLETO DO ALUNO]»")
    r.font.name = FONTE_TITULO; r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("o presente Diploma, em reconhecimento à conclusão integral dos quatro\n"
                  "níveis formativos da Escola Bíblica Epignósis — Discípulo, Crescimento,\n"
                  "Servir e Multiplicação — totalizando «[2.200–2.400]» horas de formação\n"
                  "bíblica, teológica, espiritual e ministerial.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11)
    r.font.color.rgb = COR_TEXTO

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("“Até que todos cheguemos à unidade da fé, e ao conhecimento (ἐπίγνωσιν)\n"
                  "do Filho de Deus, a homem perfeito, à medida da estatura completa de Cristo.”\n"
                  "Efésios 4.13 (ARC)")
    r.font.name = FONTE_TITULO; r.font.size = Pt(10)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    doc.add_paragraph()

    # Assinaturas (4 colunas para o diploma final)
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(tbl.rows[0].cells):
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10); r.font.color.rgb = COR_TEXTO

    rotulos_diploma = [
        ("«[Local]», «[data]»", ""),
        ("«[Director Geral]»", "Director Geral"),
        ("«[Coord. Acadêmico]»", "Coordenação Acadêmica"),
        ("«[Cons. Doutrinário]»", "Conselho Doutrinário"),
    ]
    for i, (nome, cargo) in enumerate(rotulos_diploma):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(nome)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        if cargo:
            p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(cargo)
            r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
            r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("Registo institucional: EBE-DIPL-«[ano]»-«[n.º]»   ·   Soli Deo Gloria.")
    r.font.name = FONTE_CORPO; r.font.size = Pt(8); r.font.italic = True
    r.font.color.rgb = COR_CITACAO

    out = os.path.join(os.path.dirname(__file__), "EBE-MODELOS_CERTIFICADOS.docx")
    doc.save(out)
    print("OK:", out)


if __name__ == "__main__":
    gerar()
