"""
MODELO OFICIAL DE FICHA DE MATRÍCULA — Escola Bíblica Epignósis
Formulário .docx editável e preenchível.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line


def _campo(tbl, rotulo, valor="", largura_rotulo=4.0, largura_valor=12.0):
    """Adiciona uma linha de campo (rótulo + espaço de preenchimento) numa tabela."""
    row = tbl.add_row().cells
    row[0].text = rotulo
    row[1].text = valor
    row[0].width = Cm(largura_rotulo)
    row[1].width = Cm(largura_valor)
    _shade_cell(row[0], "E8F1EC")
    for p in row[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = FONTE_TITULO
            r.font.size = Pt(10)
            r.font.color.rgb = COR_PRIMARIA
    for p in row[1].paragraphs:
        for r in p.runs:
            r.font.name = FONTE_CORPO
            r.font.size = Pt(10)


def _secao_form(doc, titulo_secao):
    """Cabeçalho de secção do formulário."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(titulo_secao.upper())
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Faz fundo
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEX_PRIMARIA)
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.2)


def _checkbox(doc, opcoes, inline=True):
    """Linha de opções com ☐ para marcar."""
    if inline:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        for i, op in enumerate(opcoes):
            r = p.add_run(f"☐  {op}    ")
            r.font.name = FONTE_CORPO
            r.font.size = Pt(10)
    else:
        for op in opcoes:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(f"☐  {op}")
            r.font.name = FONTE_CORPO
            r.font.size = Pt(10)


def _tabela_form(doc, cols=2):
    tbl = doc.add_table(rows=0, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    return tbl


def gerar():
    doc = novo_documento("Ficha de Matrícula", "EBE-FRM-001")

    # ====== CAPA / CABEÇALHO ======
    doc.add_paragraph()
    inserir_logo(doc, LOGO_PATH, largura_cm=5.5)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ἐπίγνωσις  ·  Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(9)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FICHA DE MATRÍCULA")
    r.font.name = FONTE_TITULO; r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Formulário oficial de inscrição do aluno")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    # Quadro de uso da Secretaria
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rotulos = [
        ("N.º de Matrícula", "«[      ]»"),
        ("Data de entrada", "«[__/__/____]»"),
        ("Turma / Código", "«[          ]»"),
        ("Funcionário responsável", "«[          ]»"),
    ]
    for i, (k, v) in enumerate(rotulos):
        c = tbl.rows[0].cells[i]
        c.text = ""
        _shade_cell(c, HEX_SECUNDARIA)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(k)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(v)
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("(Quadro reservado à Secretaria Acadêmica)")
    r.font.name = FONTE_CORPO; r.font.size = Pt(8)
    r.font.italic = True; r.font.color.rgb = COR_CITACAO

    # ====== 1. DADOS PESSOAIS ======
    _secao_form(doc, "1.  Dados pessoais")
    tbl = _tabela_form(doc)
    _campo(tbl, "Nome completo")
    _campo(tbl, "Data de nascimento")
    _campo(tbl, "Naturalidade")
    _campo(tbl, "Nacionalidade")
    _campo(tbl, "Sexo")
    _campo(tbl, "Estado civil")
    _campo(tbl, "Documento de identificação (BI/Passaporte)")
    _campo(tbl, "N.º do documento")
    _campo(tbl, "Filiação (Pai)")
    _campo(tbl, "Filiação (Mãe)")

    # ====== 2. CONTACTOS E MORADA ======
    _secao_form(doc, "2.  Contactos e morada")
    tbl = _tabela_form(doc)
    _campo(tbl, "Morada completa")
    _campo(tbl, "Bairro / Comuna")
    _campo(tbl, "Município / Cidade")
    _campo(tbl, "Província / Distrito")
    _campo(tbl, "País")
    _campo(tbl, "Telemóvel principal")
    _campo(tbl, "Telemóvel alternativo")
    _campo(tbl, "Endereço de e-mail")
    _campo(tbl, "Contacto WhatsApp")
    _campo(tbl, "Pessoa para emergências (nome e contacto)")

    # ====== 3. ESCOLARIDADE ======
    _secao_form(doc, "3.  Formação escolar e profissional")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("Grau de escolaridade concluído:")
    r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
    r.font.color.rgb = COR_PRIMARIA
    _checkbox(doc, [
        "Ensino Primário", "Ensino Secundário", "Ensino Médio",
        "Ensino Superior", "Pós-graduação", "Outro",
    ])

    tbl = _tabela_form(doc)
    _campo(tbl, "Última instituição frequentada")
    _campo(tbl, "Curso / Área de formação")
    _campo(tbl, "Profissão actual")
    _campo(tbl, "Entidade empregadora (se aplicável)")
    _campo(tbl, "Línguas que fala")

    # ====== 4. VIDA CRISTÃ E ECLESIAL ======
    _secao_form(doc, "4.  Vida cristã e eclesial")
    tbl = _tabela_form(doc)
    _campo(tbl, "Igreja local que frequenta")
    _campo(tbl, "Denominação")
    _campo(tbl, "Pastor responsável")
    _campo(tbl, "Contacto do pastor / liderança")
    _campo(tbl, "Tempo de conversão (anos)")
    _campo(tbl, "Data do baptismo nas águas")
    _campo(tbl, "Local do baptismo")
    _campo(tbl, "Ministério(s) em que actua")
    _campo(tbl, "Dons espirituais reconhecidos pela igreja")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("Já recebeu o baptismo no Espírito Santo?")
    r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
    r.font.color.rgb = COR_PRIMARIA
    _checkbox(doc, ["Sim", "Não", "Em busca / discernimento"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("Já participou de outros cursos bíblicos ou teológicos?")
    r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
    r.font.color.rgb = COR_PRIMARIA
    _checkbox(doc, ["Sim", "Não"])
    paragrafo(doc, "Se sim, indique quais e onde:")
    for _ in range(2):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)

    page_break(doc)

    # ====== 5. INSCRIÇÃO ACADÊMICA ======
    _secao_form(doc, "5.  Inscrição acadêmica")
    paragrafo(doc, "Nível formativo em que se inscreve (consulte o Mapa de Cursos — EBE-DOC-006):")
    _checkbox(doc, [
        "Nível 1 — Discípulo (Conhecer)",
        "Nível 2 — Crescimento (Ser)",
        "Nível 3 — Servir (Ministério)",
        "Nível 4 — Multiplicação (Reino)",
    ], inline=False)

    tbl = _tabela_form(doc)
    _campo(tbl, "Instituto")
    _campo(tbl, "Escola")
    _campo(tbl, "Programa de formação")
    _campo(tbl, "Curso(s) / Apostila(s) específicos")
    _campo(tbl, "Data prevista de início")
    _campo(tbl, "Data prevista de conclusão")

    paragrafo(doc, "Modalidade de estudo:")
    _checkbox(doc, ["Presencial", "Híbrida", "On-line"])

    paragrafo(doc, "Turno preferencial:")
    _checkbox(doc, ["Manhã", "Tarde", "Noite", "Sábado"])

    paragrafo(doc, "Disponibilidade semanal de estudo (horas):")
    _checkbox(doc, ["Até 4 h", "5 a 8 h", "9 a 15 h", "Mais de 15 h"])

    # ====== 6. MOTIVAÇÕES ======
    _secao_form(doc, "6.  Motivações para a formação")
    paragrafo(doc, "Em poucas linhas, partilhe com a Escola:")

    h3(doc, "Que o levou a procurar a Escola Bíblica Epignósis?")
    for _ in range(4):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)

    h3(doc, "Que objectivos espera alcançar com esta formação?")
    for _ in range(4):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)

    h3(doc, "Há algum chamado ministerial específico que sente da parte de Deus?")
    for _ in range(4):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)

    page_break(doc)

    # ====== 7. DECLARAÇÃO DO ALUNO ======
    _secao_form(doc, "7.  Declaração e compromisso do aluno")
    paragrafo(doc,
        "Pelo presente instrumento, o(a) candidato(a) acima identificado(a) "
        "declara, livre e voluntariamente, perante Deus e perante a Escola "
        "Bíblica Epignósis, o seguinte:")
    lista(doc, [
        "Que recebeu, leu e compreendeu a Declaração de Fé Institucional "
        "(EBE-DOC-002) da Escola Bíblica Epignósis e com ela está em "
        "comunhão e concordância de coração.",
        "Que recebeu, leu e compreendeu o Regimento Acadêmico (EBE-DOC-004) "
        "e o Projecto Pedagógico Oficial (EBE-DOC-003), comprometendo-se a "
        "cumpri-los integralmente.",
        "Que se compromete a manter conduta cristã exemplar dentro e fora "
        "do ambiente acadêmico, observando os princípios bíblicos de "
        "santidade, respeito, integridade e amor fraterno.",
        "Que se compromete a frequência, participação e dedicação às "
        "actividades propostas, cumprindo prazos e prestando as avaliações "
        "com honestidade.",
        "Que assume responsabilidade pela veracidade de todas as "
        "informações prestadas nesta ficha, ciente de que dados falsos "
        "podem implicar o cancelamento da matrícula.",
        "Que autoriza a Escola Bíblica Epignósis a tratar os seus dados "
        "pessoais exclusivamente para fins académicos, pastorais e "
        "administrativos, em conformidade com a legislação vigente.",
        "Que aceita, como referência teológica primária dos seus estudos, "
        "a Almeida Revista e Corrigida (ARC) e quaisquer outras versões "
        "indicadas pelos docentes.",
    ], ordenada=True)

    citacao(doc,
        "Estatuiu Esdras no seu coração estudar a lei do Senhor, e cumpri-la, "
        "e ensinar em Israel os seus estatutos e os seus juízos.",
        "Esdras 7.10")

    # ====== 8. ASSINATURAS ======
    _secao_form(doc, "8.  Assinaturas")
    doc.add_paragraph()
    doc.add_paragraph()

    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in tbl.rows[0].cells:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("_____________________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    rotulos = [
        ("«[Nome completo do(a) aluno(a)]»", "Assinatura do(a) aluno(a)"),
        ("«[Local], «[__]» de «[mês]» de «[ano]»", "Data"),
    ]
    for i, (nome, leg) in enumerate(rotulos):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(nome)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(leg)
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
        r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    doc.add_paragraph()

    paragrafo(doc, "")
    paragrafo(doc, "Declaração da liderança da igreja local (recomendado):", bold=True)
    paragrafo(doc,
        "Eu, abaixo assinado(a), na qualidade de líder/pastor(a) responsável "
        "pela igreja local frequentada pelo(a) candidato(a), atesto que "
        "este(a) é membro em comunhão e apoio sua matrícula na Escola "
        "Bíblica Epignósis.")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in tbl.rows[0].cells:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("_____________________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    rot2 = [
        ("«[Nome do(a) pastor(a) / líder]»", "Assinatura"),
        ("«[Igreja local]»  ·  «[Contacto]»", "Carimbo / data"),
    ]
    for i, (n, l) in enumerate(rot2):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(n)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(l)
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
        r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    page_break(doc)

    # ====== USO INTERNO DA SECRETARIA ======
    _secao_form(doc, "9.  Uso interno da Secretaria Acadêmica")

    tbl = _tabela_form(doc)
    _campo(tbl, "Documentos entregues")
    _campo(tbl, "Foto digital recebida")
    _campo(tbl, "Pré-requisitos verificados")
    _campo(tbl, "Trilha formativa atribuída")
    _campo(tbl, "Coordenador responsável")
    _campo(tbl, "Mentor designado (se aplicável)")
    _campo(tbl, "Observações pedagógicas")
    _campo(tbl, "Observações pastorais")

    paragrafo(doc, "Despacho da Secretaria:")
    _checkbox(doc, ["Matrícula deferida", "Matrícula condicional", "Indeferida"])

    paragrafo(doc, "Justificação (se aplicável):")
    for _ in range(3):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)

    doc.add_paragraph()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in tbl.rows[0].cells:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("_____________________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    rot3 = [
        ("«[Secretaria Acadêmica]»", "Assinatura e carimbo"),
        ("«[Coordenação Acadêmica]»", "Aprovação"),
    ]
    for i, (n, l) in enumerate(rot3):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(n)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(l)
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
        r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    selo_final(doc)

    out = os.path.join(os.path.dirname(__file__), "EBE-FRM-001_Ficha_de_Matricula.docx")
    doc.save(out)
    print("OK:", out)


if __name__ == "__main__":
    gerar()
