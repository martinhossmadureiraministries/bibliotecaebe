"""
Módulo de estilos comum para os documentos institucionais da
Escola Bíblica Epignósis (EBE).

Design: Académico formal (pt-PT, ARC).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# === Paleta institucional (harmonizada com o logotipo EBE) ===
COR_PRIMARIA   = RGBColor(0x1B, 0x3A, 0x5C)   # azul-marinho do logo
COR_SECUNDARIA = RGBColor(0x2E, 0x7D, 0x4F)   # verde institucional do logo
COR_TERCIARIA  = RGBColor(0xC9, 0xA1, 0x4B)   # dourado para selos
COR_TEXTO      = RGBColor(0x1A, 0x1A, 0x1A)   # quase preto
COR_LINHA      = RGBColor(0xB8, 0xB8, 0xB8)
COR_CITACAO    = RGBColor(0x55, 0x55, 0x55)

# Códigos HEX (string, para tabelas e bordas)
HEX_PRIMARIA   = "1B3A5C"
HEX_SECUNDARIA = "2E7D4F"
HEX_TERCIARIA  = "C9A14B"

FONTE_TITULO = "Garamond"
FONTE_CORPO  = "Garamond"
FONTE_MONO   = "Consolas"


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), kwargs[edge].get("val", "single"))
            border.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            border.set(qn("w:color"), kwargs[edge].get("color", "B8B8B8"))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_line(paragraph, color=HEX_PRIMARIA, size=8):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_page(section, top=2.5, bottom=2.5, left=3.0, right=2.5):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(6)  # WD_BREAK.PAGE = 7? use OxmlElement for safety
    # fallback
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def configurar_estilos_base(doc: Document):
    """Aplica estilos académicos padrão (pt-PT)."""
    # Estilo Normal
    style = doc.styles["Normal"]
    style.font.name = FONTE_CORPO
    style.font.size = Pt(12)
    style.font.color.rgb = COR_TEXTO
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.4
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Margens (A4 académico)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        _set_page(section)


def add_capa(doc, supratitulo, titulo, subtitulo=None, codigo=None, ano=None, com_logo=True):
    """Capa académica sóbria com logotipo."""
    # Espaço superior
    doc.add_paragraph()

    # Logotipo
    if com_logo and _os.path.exists(LOGO_PATH):
        inserir_logo(doc, LOGO_PATH, largura_cm=6.5)
    else:
        for _ in range(3):
            doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ἐπίγνωσις  ·  Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=8)

    for _ in range(3):
        doc.add_paragraph()

    if supratitulo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(supratitulo.upper())
        r.font.name = FONTE_TITULO
        r.font.size = Pt(11)
        r.font.color.rgb = COR_SECUNDARIA
        r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.name = FONTE_TITULO
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    if subtitulo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitulo)
        r.font.name = FONTE_TITULO
        r.font.size = Pt(13)
        r.font.italic = True
        r.font.color.rgb = COR_TEXTO

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)

    if codigo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Documento institucional · {codigo}")
        r.font.name = FONTE_CORPO
        r.font.size = Pt(10)
        r.font.color.rgb = COR_CITACAO

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Edição oficial · {ano or '2026'}")
    r.font.name = FONTE_CORPO
    r.font.size = Pt(10)
    r.font.color.rgb = COR_CITACAO

    page_break(doc)


def add_marco_filosofico(doc):
    """Página do marco filosófico — aparece após a capa."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MARCO FILOSÓFICO")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.right_indent = Cm(2)
    r = p.add_run(
        "“Acreditamos que o verdadeiro conhecimento de Deus transforma "
        "a mente pela verdade das Escrituras, o coração pela acção do "
        "Espírito Santo e a vida pelo compromisso de viver e anunciar "
        "o Evangelho de Jesus Cristo.”"
    )
    r.font.name = FONTE_TITULO
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = COR_PRIMARIA

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Escola Bíblica Epignósis —")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(10)
    r.font.color.rgb = COR_CITACAO

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "“Até que todos cheguemos à unidade da fé e ao "
        "pleno conhecimento (ἐπίγνωσις) do Filho de Deus, "
        "a homem perfeito, à medida da estatura completa de Cristo.”\nEfésios 4.13"
    )
    r.font.name = FONTE_CORPO
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COR_CITACAO

    page_break(doc)


def h1(doc, texto, numero=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    txt = f"{numero}. {texto}" if numero else texto
    r = p.add_run(txt.upper())
    r.font.name = FONTE_TITULO
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA
    _add_horizontal_line(p, color=HEX_PRIMARIA, size=6)
    return p


def h2(doc, texto, numero=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    txt = f"{numero}. {texto}" if numero else texto
    r = p.add_run(txt)
    r.font.name = FONTE_TITULO
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA
    return p


def h3(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(texto)
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA
    return p


def paragrafo(doc, texto, italic=False, bold=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.name = FONTE_CORPO
    r.font.size = Pt(12)
    r.font.italic = italic
    r.font.bold = bold
    return p


def artigo(doc, numero, texto, paragrafo_unico=None, paragrafos=None, incisos=None):
    """
    Renderiza um artigo no estilo jurídico/académico:

      Art. Nº — texto principal.
        § 1.º texto…
        § 2.º texto…
        I — inciso…
        II — inciso…
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"Art. {numero}.º  ")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA
    r2 = p.add_run(texto)
    r2.font.name = FONTE_CORPO
    r2.font.size = Pt(12)

    if paragrafo_unico:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run("Parágrafo único. ")
        r.font.bold = True
        r.font.name = FONTE_TITULO
        r.font.size = Pt(11.5)
        r2 = p.add_run(paragrafo_unico)
        r2.font.name = FONTE_CORPO
        r2.font.size = Pt(11.5)

    if paragrafos:
        for i, t in enumerate(paragrafos, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(f"§ {i}.º  ")
            r.font.bold = True
            r.font.name = FONTE_TITULO
            r.font.size = Pt(11.5)
            r2 = p.add_run(t)
            r2.font.name = FONTE_CORPO
            r2.font.size = Pt(11.5)

    if incisos:
        romanos = ["I","II","III","IV","V","VI","VII","VIII","IX","X","XI","XII","XIII","XIV","XV","XVI","XVII","XVIII","XIX","XX","XXI","XXII","XXIII","XXIV","XXV"]
        for i, t in enumerate(incisos):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(f"{romanos[i]} — ")
            r.font.bold = True
            r.font.name = FONTE_TITULO
            r.font.size = Pt(11.5)
            r2 = p.add_run(t)
            r2.font.name = FONTE_CORPO
            r2.font.size = Pt(11.5)


def capitulo(doc, numero_romano, titulo):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"CAPÍTULO {numero_romano}")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(titulo.upper())
    r.font.name = FONTE_TITULO
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_PRIMARIA, size=4)


def secao_titulada(doc, titulo):
    """Título de secção (sem numeração de artigo)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(titulo)
    r.font.name = FONTE_TITULO
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = COR_SECUNDARIA


def citacao(doc, texto, referencia=None):
    """Bloco de citação bíblica recuado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f"“{texto}”")
    r.font.name = FONTE_CORPO
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = COR_CITACAO
    if referencia:
        r2 = p.add_run(f"  ({referencia}, ARC)")
        r2.font.name = FONTE_CORPO
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = COR_SECUNDARIA


def lista(doc, itens, ordenada=False):
    for i, item in enumerate(itens, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(2)
        marca = f"{i}. " if ordenada else "•  "
        r = p.add_run(marca)
        r.font.name = FONTE_TITULO
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = COR_SECUNDARIA
        r2 = p.add_run(item)
        r2.font.name = FONTE_CORPO
        r2.font.size = Pt(12)


def selo_final(doc):
    """Selo institucional de encerramento."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    _add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Soli Deo Gloria")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = COR_CITACAO


def cabecalho_rodape(doc, titulo_doc, codigo_doc):
    """Cabeçalho discreto + rodapé com paginação."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in doc.sections:
        section.different_first_page_header_footer = True

        # Cabeçalho
        header = section.header
        if header.paragraphs:
            ph = header.paragraphs[0]
        else:
            ph = header.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ph.text = ""
        r = ph.add_run(f"Escola Bíblica Epignósis  ·  {titulo_doc}")
        r.font.name = FONTE_TITULO
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = COR_SECUNDARIA

        # Rodapé com número de página
        footer = section.footer
        if footer.paragraphs:
            pf = footer.paragraphs[0]
        else:
            pf = footer.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pf.add_run(f"{codigo_doc}  ·  ")
        r.font.name = FONTE_CORPO
        r.font.size = Pt(9)
        r.font.color.rgb = COR_CITACAO

        # campo PAGE
        run = pf.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = FONTE_CORPO
        run.font.size = Pt(9)
        run.font.color.rgb = COR_CITACAO


import os as _os

ASSETS_DIR = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "_assets")
LOGO_PATH        = _os.path.join(ASSETS_DIR, "logo_ebe.png")
LOGO_PEQUENO     = _os.path.join(ASSETS_DIR, "logo_pequeno.png")
LOGO_EMBLEMA     = _os.path.join(ASSETS_DIR, "logo_emblema.png")
LOGO_MARCA_DAGUA = _os.path.join(ASSETS_DIR, "logo_marca_dagua.png")


def inserir_logo(doc, caminho=None, largura_cm=5.5, alinhamento="center"):
    """Insere o logotipo num parágrafo centralizado."""
    caminho = caminho or LOGO_PATH
    p = doc.add_paragraph()
    if alinhamento == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alinhamento == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    r.add_picture(caminho, width=Cm(largura_cm))
    return p


def novo_documento(titulo_doc, codigo_doc):
    """Cria um Document padronizado já com estilos e cabeçalho/rodapé."""
    doc = Document()
    configurar_estilos_base(doc)
    cabecalho_rodape(doc, titulo_doc, codigo_doc)
    return doc
