"""Gera o compêndio unificado por merge dos 8 docs já gerados."""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer


HERE = os.path.dirname(os.path.abspath(__file__))


def build_capa():
    """Cria a capa institucional + sumário geral num doc base."""
    doc = novo_documento("Compêndio Institucional", "EBE-COMPÊNDIO")

    # CAPA
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMPÊNDIO INSTITUCIONAL OFICIAL")
    r.font.name = FONTE_TITULO; r.font.size = Pt(13)
    r.font.color.rgb = COR_SECUNDARIA; r.font.bold = True

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ἐπίγνωσις")
    r.font.name = FONTE_TITULO; r.font.size = Pt(22)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color="2E7D4F", size=8)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conhecer a Deus.   Viver a Palavra.   Manifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(14)
    r.font.italic = True; r.font.color.rgb = COR_PRIMARIA

    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Reunião oficial dos oito documentos institucionais\n(EBE-DOC-001 a EBE-DOC-008)")
    r.font.name = FONTE_CORPO; r.font.size = Pt(11); r.font.color.rgb = COR_TEXTO

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); _add_horizontal_line(p, color="2E7D4F", size=4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Edição oficial · 2026")
    r.font.name = FONTE_CORPO; r.font.size = Pt(10); r.font.color.rgb = COR_CITACAO

    page_break(doc)

    # MARCO FILOSÓFICO
    add_marco_filosofico(doc)

    # APRESENTAÇÃO
    h1(doc, "Apresentação do Compêndio")
    paragrafo(doc,
        "O presente Compêndio reúne, em um único volume oficial, os oito "
        "documentos institucionais que estruturam a vida acadêmica, "
        "espiritual, doutrinária e ministerial da Escola Bíblica Epignósis. "
        "Estes documentos são complementares entre si e devem ser lidos em "
        "conjunto para uma compreensão integral da identidade, da filosofia "
        "educativa e da arquitectura curricular da instituição.")
    paragrafo(doc,
        "Cada documento conserva numeração e paginação próprias, podendo "
        "ser destacado, impresso e consultado de forma autónoma. Na "
        "Secretaria Acadêmica, as edições oficiais individuais (formato "
        ".docx editável) constituem a fonte primária de consulta.")

    # SUMÁRIO GERAL
    h1(doc, "Sumário Geral")
    docs = [
        ("EBE-DOC-001", "Identidade Institucional",          "Missão, Visão, Valores e Filosofia de Ensino."),
        ("EBE-DOC-002", "Declaração de Fé Institucional",    "Os 16 artigos doutrinários fundamentais."),
        ("EBE-DOC-003", "Projecto Pedagógico Oficial (PPO)", "Bases acadêmicas, pedagógicas e ministeriais."),
        ("EBE-DOC-004", "Regimento Acadêmico",               "Normas de organização e funcionamento."),
        ("EBE-DOC-005", "Arquitectura Oficial",              "Estrutura hierárquica de sete níveis."),
        ("EBE-DOC-006", "Mapa Oficial de Cursos",            "Institutos, Escolas e Trilha formativa."),
        ("EBE-DOC-007", "Duração Oficial dos Cursos",        "Carga horária por nível, instituto e escola."),
        ("EBE-DOC-008", "Sistema de Pré-Requisitos",         "Princípios de progressão académica e espiritual."),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(["Código", "Documento", "Conteúdo"]):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _shade_cell(hdr[i], "1B3A5C")
    for cod, t, conteudo in docs:
        row = tbl.add_row().cells
        row[0].text = cod; row[1].text = t; row[2].text = conteudo
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = FONTE_CORPO; run.font.size = Pt(11)

    paragrafo(doc, "")
    paragrafo(doc,
        "A interpretação e aplicação dos documentos seguem o princípio "
        "hermenêutico da unidade institucional: todos os documentos devem "
        "ser lidos à luz da Declaração de Fé (EBE-DOC-002) e do Projecto "
        "Pedagógico Oficial (EBE-DOC-003).")

    # Página de divisão antes do bloco geral
    page_break(doc)

    out_capa = os.path.join(HERE, "_capa_compendio_tmp.docx")
    doc.save(out_capa)
    return out_capa


def build_compendio():
    capa = build_capa()
    composer = Composer(Document(capa))
    for nome in [
        "EBE-DOC-001_Identidade_Institucional.docx",
        "EBE-DOC-002_Declaracao_de_Fe.docx",
        "EBE-DOC-003_Projecto_Pedagogico_Oficial.docx",
        "EBE-DOC-004_Regimento_Academico.docx",
        "EBE-DOC-005_Arquitectura_Oficial.docx",
        "EBE-DOC-006_Mapa_de_Cursos.docx",
        "EBE-DOC-007_Duracao_Oficial.docx",
        "EBE-DOC-008_Sistema_de_Pre_Requisitos.docx",
    ]:
        path = os.path.join(HERE, nome)
        composer.append(Document(path))

    final = os.path.join(HERE, "EBE_COMPENDIO_INSTITUCIONAL_OFICIAL.docx")
    composer.save(final)
    os.remove(capa)
    print("OK:", final)


if __name__ == "__main__":
    build_compendio()
