"""
MODELO OFICIAL DE FICHA DE AVALIAÇÃO DO ALUNO — Escola Bíblica Epignósis

Avaliação integral em quatro dimensões (Conhecer, Crer, Viver, Servir),
com notas, observações pedagógicas e pastorais.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _estilos import *
from _estilos import _shade_cell, _add_horizontal_line


def _campo(tbl, rotulo, valor=""):
    row = tbl.add_row().cells
    row[0].text = rotulo; row[1].text = valor
    _shade_cell(row[0], "E8F1EC")
    for p in row[0].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
            r.font.color.rgb = COR_PRIMARIA
    for p in row[1].paragraphs:
        for r in p.runs:
            r.font.name = FONTE_CORPO; r.font.size = Pt(10)


def _secao_form(doc, titulo_secao):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"  {titulo_secao.upper()}")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEX_PRIMARIA)
    pPr.append(shd)


def _linhas(doc, n=3):
    for _ in range(n):
        p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)


def _tabela_notas(doc, instrumentos):
    """Tabela 5 col: Instrumento | Peso (%) | Nota (0–10) | Ponderada | Observações."""
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    headers = ["Instrumento de avaliação", "Peso (%)", "Nota (0–10)", "Nota ponderada", "Observações"]
    widths  = [5.5, 1.8, 2.0, 2.4, 4.5]
    for i, t in enumerate(headers):
        hdr[i].text = ""
        hdr[i].width = Cm(widths[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for inst in instrumentos:
        row = tbl.add_row().cells
        for i, c in enumerate(row):
            c.width = Cm(widths[i])
        row[0].text = inst
        for j in (1, 2, 3, 4):
            row[j].text = "«[…]»"
        for c in row:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(9)
    return tbl


def gerar():
    doc = novo_documento("Ficha de Avaliação", "EBE-FRM-AVA")

    # CAPA
    doc.add_paragraph()
    inserir_logo(doc, LOGO_PATH, largura_cm=5.0)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ἐπίγνωσις · Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO; r.font.size = Pt(9)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph(); _add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("FICHA DE AVALIAÇÃO DO(A) ALUNO(A)")
    r.font.name = FONTE_TITULO; r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Avaliação integral nas quatro dimensões — Conhecer · Crer · Viver · Servir")
    r.font.name = FONTE_TITULO; r.font.size = Pt(11)
    r.font.italic = True; r.font.color.rgb = COR_SECUNDARIA

    # 1. IDENTIFICAÇÃO
    _secao_form(doc, "1.  Identificação")
    tbl = doc.add_table(rows=0, cols=2)
    _campo(tbl, "Nome do(a) aluno(a)", "«[…]»")
    _campo(tbl, "N.º de matrícula", "«[…]»")
    _campo(tbl, "Igreja local", "«[…]»")
    _campo(tbl, "Instituto · Escola · Curso", "«[…]»")
    _campo(tbl, "Módulo / Apostila avaliada", "«[…]»")
    _campo(tbl, "Docente responsável", "«[…]»")
    _campo(tbl, "Período avaliativo", "«[__/__/____]» a «[__/__/____]»")
    _campo(tbl, "Tipo de avaliação", "«[Formativa contínua / Final de módulo / Final de curso]»")
    _campo(tbl, "Carga horária ministrada", "«[…] horas»")
    _campo(tbl, "Frequência do(a) aluno(a)", "«[…] %»")

    # 2. CRITÉRIO E ESCALA
    _secao_form(doc, "2.  Critério de avaliação e escala institucional")
    paragrafo(doc,
        "A Escola Bíblica Epignósis adopta avaliação formativa contínua, "
        "integrando conhecimento, convicção, conduta e capacidade ministerial. "
        "Cada dimensão é pontuada na escala de 0 a 10. A nota final do "
        "período resulta da média ponderada dos instrumentos aplicados.")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Faixa", "Conceito", "Significado pedagógico"]
    widths  = [2.5, 4.0, 9.5]
    hdr = tbl.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].text = ""; hdr[i].width = Cm(widths[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    escala = [
        ("9,0 – 10,0", "Excelente",   "Domínio sólido, aplicação madura e testemunho exemplar."),
        ("7,5 – 8,9",  "Muito bom",   "Boa compreensão e aplicação consistente do conteúdo."),
        ("6,0 – 7,4",  "Satisfatório","Compreensão básica adequada, com pontos a aprofundar."),
        ("5,0 – 5,9",  "Suficiente",  "Aprovação com necessidade de reforço pedagógico."),
        ("0,0 – 4,9",  "Insuficiente","Reprovado(a) — recomendado refazer a etapa avaliada."),
    ]
    for r_data in escala:
        row = tbl.add_row().cells
        for i, v in enumerate(r_data):
            row[i].text = v; row[i].width = Cm(widths[i])
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    paragrafo(doc, "Média mínima institucional para aprovação: 6,0 (seis vírgula zero), conforme Regimento Acadêmico (EBE-DOC-004, Art. 16.º).", italic=True)

    page_break(doc)

    # 3. AVALIAÇÃO COGNITIVA — CONHECER
    _secao_form(doc, "3.  Dimensão CONHECER — Avaliação cognitiva")
    paragrafo(doc, "Verifica o domínio do conteúdo bíblico-teológico estudado, a capacidade de interpretação, de articulação doutrinária e de aplicação hermenêutica.")
    _tabela_notas(doc, [
        "Prova escrita",
        "Questionário objectivo",
        "Trabalho académico / pesquisa",
        "Estudo bíblico elaborado",
        "Apresentação oral",
    ])
    paragrafo(doc, "Subtotal da dimensão CONHECER:")
    p = doc.add_paragraph(); _add_horizontal_line(p, color="C8C8C8", size=4)

    # 4. AVALIAÇÃO DE CONVICÇÃO — CRER
    _secao_form(doc, "4.  Dimensão CRER — Convicção e fundamentação da fé")
    paragrafo(doc, "Verifica se o(a) aluno(a) interioriza biblicamente o conteúdo, fundamenta as suas convicções nas Escrituras e mantém adesão à Declaração de Fé Institucional.")
    h3(doc, "Quesitos qualitativos (escala 0–10)")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Quesito", "Nota", "Observações"]
    widths  = [8.0, 2.0, 6.0]
    hdr = tbl.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].text = ""; hdr[i].width = Cm(widths[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    quesitos_crer = [
        "Articulação bíblica das convicções pessoais.",
        "Adesão à Declaração de Fé Institucional.",
        "Capacidade de defender a fé com mansidão (1 Pe 3.15).",
        "Discernimento espiritual e doutrinário.",
    ]
    for q in quesitos_crer:
        row = tbl.add_row().cells
        for i, c in enumerate(row):
            c.width = Cm(widths[i])
        row[0].text = q; row[1].text = "«[…]»"; row[2].text = "«[…]»"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    # 5. AVALIAÇÃO DE VIDA — VIVER
    _secao_form(doc, "5.  Dimensão VIVER — Conduta cristã e maturidade")
    paragrafo(doc, "Verifica a coerência entre o conteúdo estudado e a vida diária do(a) aluno(a), reflectida no carácter, na convivência e no testemunho público.")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, t in enumerate(["Quesito", "Nota", "Observações"]):
        hdr[i].text = ""; hdr[i].width = Cm(widths[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    quesitos_viver = [
        "Vida devocional consistente (oração e leitura bíblica).",
        "Humildade e disposição para aprender.",
        "Integridade e ética cristã.",
        "Respeito, amor e bom relacionamento com colegas e docentes.",
        "Pontualidade, assiduidade e responsabilidade.",
        "Coerência entre o que estuda e o que pratica (Tg 1.22).",
    ]
    for q in quesitos_viver:
        row = tbl.add_row().cells
        for i, c in enumerate(row): c.width = Cm(widths[i])
        row[0].text = q; row[1].text = "«[…]»"; row[2].text = "«[…]»"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    page_break(doc)

    # 6. AVALIAÇÃO MINISTERIAL — SERVIR
    _secao_form(doc, "6.  Dimensão SERVIR — Capacidade ministerial e prática")
    paragrafo(doc, "Verifica a aplicação dos dons e dos conhecimentos no serviço cristão, na igreja local e na vida prática.")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, t in enumerate(["Quesito", "Nota", "Observações"]):
        hdr[i].text = ""; hdr[i].width = Cm(widths[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    quesitos_servir = [
        "Participação activa nas actividades práticas.",
        "Disposição para servir e para o trabalho em equipa.",
        "Capacidade de ensinar e discipular outros.",
        "Exercício equilibrado de dons espirituais.",
        "Iniciativa para projectos ministeriais ou missionários.",
        "Frutos observáveis na igreja local.",
    ]
    for q in quesitos_servir:
        row = tbl.add_row().cells
        for i, c in enumerate(row): c.width = Cm(widths[i])
        row[0].text = q; row[1].text = "«[…]»"; row[2].text = "«[…]»"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    # 7. CONSOLIDAÇÃO
    _secao_form(doc, "7.  Consolidação final")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Dimensão", "Peso (%)", "Nota (0–10)", "Nota ponderada"]
    widths_c  = [5.0, 3.0, 3.0, 4.0]
    hdr = tbl.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].text = ""; hdr[i].width = Cm(widths_c[i])
        _shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    dims = [
        ("Conhecer", "40 %", "«[…]»", "«[…]»"),
        ("Crer",     "20 %", "«[…]»", "«[…]»"),
        ("Viver",    "20 %", "«[…]»", "«[…]»"),
        ("Servir",   "20 %", "«[…]»", "«[…]»"),
    ]
    for d in dims:
        row = tbl.add_row().cells
        for i, v in enumerate(d):
            row[i].text = v; row[i].width = Cm(widths_c[i])
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = FONTE_CORPO; r.font.size = Pt(10)
    # linha final
    row = tbl.add_row().cells
    finais = ["MÉDIA FINAL", "100 %", "", "«[…]»"]
    for i, v in enumerate(finais):
        row[i].text = v; row[i].width = Cm(widths_c[i])
        _shade_cell(row[i], HEX_SECUNDARIA)
        for p in row[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.name = FONTE_TITULO; r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    paragrafo(doc, "Situação final:")
    for s in ["Aprovado(a)", "Aprovado(a) com louvor (≥ 9,0)", "Aprovado(a) por compensação", "Reprovado(a)", "Pendente — reavaliação"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f"☐  {s}")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)

    # 8. PARECERES
    _secao_form(doc, "8.  Parecer pedagógico")
    paragrafo(doc, "Avaliação qualitativa do(a) docente sobre o aproveitamento, dificuldades e potencialidades do(a) aluno(a):")
    _linhas(doc, 5)
    paragrafo(doc, "Recomendações pedagógicas (estudos complementares, reforço, mentoria, etc.):")
    _linhas(doc, 3)

    _secao_form(doc, "9.  Parecer pastoral")
    paragrafo(doc, "Observações sobre a vida espiritual, o carácter, a maturidade e o testemunho do(a) aluno(a):")
    _linhas(doc, 5)
    paragrafo(doc, "Encaminhamento pastoral (se aplicável):")
    _linhas(doc, 3)

    citacao(doc,
        "Procura apresentar-te a Deus aprovado, como obreiro que não tem de "
        "que se envergonhar, que maneja bem a palavra da verdade.",
        "2 Timóteo 2.15")

    # 10. CIÊNCIA DO ALUNO
    _secao_form(doc, "10.  Ciência do(a) aluno(a)")
    paragrafo(doc,
        "Declaro que tomei conhecimento da presente avaliação e que dela "
        "recebi cópia. Comprometo-me a empenhar-me na superação dos "
        "pontos indicados, na busca da maturidade integral em Cristo.")

    doc.add_paragraph()
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in tbl.rows[0].cells:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("__________________________________")
        r.font.name = FONTE_CORPO; r.font.size = Pt(10)
    rot = [
        ("«[Aluno(a)]»", "Ciência e concordância"),
        ("«[Docente responsável]»", "Avaliação"),
        ("«[Coordenação Acadêmica]»", "Visto institucional"),
    ]
    for i, (n, l) in enumerate(rot):
        c = tbl.rows[1].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(n)
        r.font.name = FONTE_TITULO; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = COR_PRIMARIA
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(l)
        r2.font.name = FONTE_CORPO; r2.font.size = Pt(9)
        r2.font.italic = True; r2.font.color.rgb = COR_SECUNDARIA

    selo_final(doc)

    out = os.path.join(os.path.dirname(__file__), "EBE-FRM-004_Ficha_de_Avaliacao.docx")
    doc.save(out); print("OK:", out)


if __name__ == "__main__":
    gerar()
